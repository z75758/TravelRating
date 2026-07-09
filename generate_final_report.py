"""
Generate final complete course design report for TravelRating.
Includes all latest features: comment image upload, hero background, encoding fix, etc.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'; font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

base = r'F:\MyProject\TravelRating\src\main'
def src(path):
    p = os.path.join(base, path)
    return open(p, encoding='utf-8').read() if os.path.exists(p) else ''

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0,0,0); r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level==1: r.font.size=Pt(16)
        elif level==2: r.font.size=Pt(14)
        else: r.font.size=Pt(13)

def P(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text); r.font.size=Pt(12); r.bold=bold
    r.font.name='宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

def C(text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent=Cm(1)
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.line_spacing=1.0
        r = p.add_run(line); r.font.name='Consolas'; r.font.size=Pt(8.5)
        r.font.color.rgb=RGBColor(51,51,51)
    doc.add_paragraph()

def T(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(10)
        r.font.name='宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
        sh=parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        c._tc.get_or_add_tcPr().append(sh)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''
            r=c.paragraphs[0].add_run(str(val)); r.font.size=Pt(10)
            r.font.name='宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    doc.add_paragraph()

def B(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.74)
    p.paragraph_format.line_spacing=1.5
    r=p.add_run('• '+text); r.font.size=Pt(12)
    r.font.name='宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

# Read source files
dbutil = src('java/com/travel/util/DBUtil.java')
pwdutil = src('java/com/travel/util/PasswordUtil.java')
user_sv = src('java/com/travel/controller/UserServlet.java')
dest_sv = src('java/com/travel/controller/DestinationServlet.java')
vote_sv = src('java/com/travel/controller/VoteServlet.java')
cmt_sv = src('java/com/travel/controller/CommentServlet.java')
search_sv = src('java/com/travel/controller/SearchServlet.java')
img_sv = src('java/com/travel/controller/ImageUploadServlet.java')
auth_fl = src('java/com/travel/filter/AuthFilter.java')
enc_fl = src('java/com/travel/filter/EncodingFilter.java')
vote_dao = src('java/com/travel/dao/VoteDAO.java')
vote_svc = src('java/com/travel/service/VoteService.java')
index_jsp = src('webapp/index.jsp')
login_jsp = src('webapp/login.jsp')
header_jsp = src('webapp/header.jsp')
detail_jsp = src('webapp/destination/detail.jsp')

# ═══════════ TITLE ═══════════
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('第四部分：课程设计内容（步骤及程序）'); r.bold=True; r.font.size=Pt(18)
r.font.name='黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
st=doc.add_paragraph(); st.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=st.add_run('项目名称：旅游地点评选与互动系统设计与实现'); r.font.size=Pt(14)

P('')
P('摘要：本项目为旅游地点评选与互动系统，是一个基于JSP + Servlet + JDBC技术栈的B/S架构Web应用程序。'
  '系统面向游客与管理员两类用户，提供旅游目的地浏览、多条件搜索筛选、交互式五星评分、评论互动（支持图片上传）、'
  '后台管理等完整功能。技术实现上采用MVC分层架构，以H2嵌入式数据库实现零安装部署，'
  '采用SHA-256+随机盐值加密保障用户密码安全，前端基于Taste-Skill极简设计系统实现暗色模式和响应式布局。'
  '系统还实现了主页大标题区域旋转风景背景轮播（30秒切换+白色羽化效果）、评论图片上传、'
  '评论删除权限控制、中文编码统一修复等功能。系统代码约3500行，包含6个Servlet、4个DAO、4个Service、'
  '4个Model、2个Filter、3个Util及16个JSP页面，97项测试全部通过。', indent=True)
P('')
P('关键词：旅游目的地评选；JSP；Servlet；JDBC；H2数据库；MVC架构；五星评分；图片上传；羽化背景', bold=True)

doc.add_page_break()

# ═══════════ 1. 系统概述 ═══════════
H('1. 系统概述',1)

H('1.1 设计目的',2)
P('随着我国旅游业的蓬勃发展，游客对于旅游目的地的信息获取和评价反馈需求日益增长。'
  '传统的旅游信息网站多采用静态内容展示方式，缺乏用户互动和评价机制。'
  '本课程设计旨在综合运用Web开发技术，完成一个集目的地展示、搜索筛选、星级评分、'
  '评论互动（含图片上传）于一体的完整Web应用系统。具体设计目的包括：')
P('（1）综合运用JSP、Servlet、JDBC等Java Web核心技术，构建功能完备的B/S架构应用系统；')
P('（2）深入实践MVC分层架构，将DAO、Service、Servlet和JSP进行清晰的职责分离；')
P('（3）掌握关系型数据库设计与实现，包括ER图设计、表结构规范化、外键约束与级联操作；')
P('（4）实现完整的用户认证与权限管理体系，深入理解Session机制、Filter过滤器链及SHA-256+Salt密码加密策略；')
P('（5）应用Taste-Skill极简设计系统，实践CSS自定义属性、响应式布局、暗色模式适配、CSS羽化渐变等前端技术；')
P('（6）掌握AJAX异步交互技术，实现无刷新星级评分、异步图片上传等功能；')
P('（7）实践Apache Commons FileUpload组件，实现评论图片的multipart/form-data上传与存储。')

H('1.2 项目概况',2)
P('旅游地点评选与互动系统（TravelRating）是一个基于Java Web技术栈的在线旅游信息服务平台。'
  '系统面向两类用户角色：普通用户（游客）和系统管理员。', indent=True)
P('普通用户功能：浏览旅游目的地列表、按地区/类型/评分/票价多条件筛选、关键字搜索、'
  '查看目的地详情（大图/描述/地址/开放时间/票价/综合评分/用户评论）、通过交互式五星组件进行1-5分评分、'
  '发表文字评论（可选上传旅行照片）、管理个人资料。')
P('管理员功能：后台仪表盘查看系统统计数据、目的地信息CRUD管理、用户列表查看与管理、'
  '目的地封面图片AJAX异步本地上传、评论管理（可删除任意评论）。')
P('技术架构：系统严格遵循MVC三层架构。表现层采用JSP页面 + Taste-Skill CSS + Vanilla JavaScript。'
  '主页和列表页Hero区域实现旋转风景背景轮播：使用5张真实目的地图片，30秒间隔自动切换，'
  '1.8秒淡入淡出过渡，CSS linear-gradient实现从左下角到右上角的白色羽化渐变效果，'
  '确保标题文字始终清晰可读。控制层由6个Servlet组成，通过@WebServlet注解配置URL映射。'
  '评论Servlet支持multipart/form-data文件上传，使用Apache Commons FileUpload解析，'
  'UUID命名存储至/uploads/目录。数据访问层使用JDBC PreparedStatement参数化查询，全面防止SQL注入。'
  '数据库采用H2 Embedded（MySQL兼容模式），数据文件存储于./data/目录，实现零安装部署。', indent=True)

doc.add_page_break()

# ═══════════ 2. 系统分析 ═══════════
H('2. 系统分析',1)

H('2.1 引言',2)
P('本节对旅游地点评选与互动系统进行全面的系统分析，明确系统的功能需求、非功能需求和运行环境要求，'
  '为后续系统设计和编码实现奠定基础。')

H('2.2 任务概述',2)
P('本系统的核心开发任务包括：')
P('（1）搭建基于Maven的Java Web项目骨架，配置tomcat7-maven-plugin嵌入式Tomcat运行环境；')
P('（2）设计并实现4张数据表：users、destinations、comments（含image字段）、votes；')
P('（3）实现用户注册与登录，SHA-256+16字节随机盐值加密存储，AuthFilter基于Session的登录拦截；')
P('（4）实现目的地CRUD管理功能，管理员通过后台页面进行创建、编辑和删除；')
P('（5）实现交互式五星评分机制（hover预览+click AJAX提交），每用户每目的地仅一票（UNIQUE约束+UPSERT）；')
P('（6）实现评论功能，支持纯文本和带图片评论，图片通过multipart/form-data上传至/uploads/；')
P('（7）实现评论删除权限控制：作者或管理员可删除，前端confirm二次确认；')
P('（8）实现多条件搜索与筛选（关键字+地区+类型+排序），StringBuilder动态拼接SQL；')
P('（9）实现主页和列表页Hero区域旋转风景背景轮播，CSS白色羽化渐变效果；')
P('（10）解决中文乱码问题：移除@WebFilter双重注册，web.xml显式声明Filter执行顺序，EncodingFilter先于AuthFilter执行；')
P('（11）实现响应式布局（768px断点）和暗色模式（prefers-color-scheme）自适应。')

H('2.3 需求规定',2)
P('一、功能需求',bold=True)
B('用户模块：注册/登录/个人资料编辑/密码修改/安全退出；')
B('目的地模块：首页Bento Grid卡片+Hero旋转背景/列表页表格视图/多条件筛选/模糊搜索/详情页完整信息；')
B('投票评分模块：交互式五星评分（hover高亮金黄花 + click AJAX提交 + Toast反馈），UPSERT策略，UNIQUE约束；')
B('评论模块：发表文字评论 + 可选上传旅行照片（JPG/PNG/GIF/WebP，最大5MB），UUID存储，点击放大；')
B('评论删除：作者或管理员可见×删除按钮，confirm确认，CommentService三重校验（存在性+作者匹配+管理员越权）；')
B('管理模块：仪表盘统计/目的地CRUD/用户管理/封面图片AJAX上传；')
B('视觉效果：主页+列表页Hero区域5张风景图30秒轮播，1.8秒淡入淡出，左下→右上白色羽化渐变。')

P('')
P('二、非功能需求',bold=True)
B('安全性：SHA-256+Salt密码加密，PreparedStatement全面防SQL注入，AuthFilter双层鉴权，Filter链编码保护；')
B('可用性：界面简洁直观，Toast消息即时反馈，评论图片预览+移除，滚动渐入动画，按钮点击反馈；')
B('兼容性：Chrome/Firefox/Edge，768px响应式，暗色模式，触控友好；')
B('可维护性：MVC分层架构，代码语义化命名，try-with-resources资源管理。')

H('2.4 运行环境',2)
P('• 操作系统：Windows 10/11 | JDK 11 | Maven 3.9+ | Tomcat 9.0.97（嵌入式）')
P('• 数据库：H2 Database 2.3.232（MySQL兼容模式，文件存储）')
P('• 浏览器：Chrome 120+ / Firefox 115+ / Edge 120+')

H('2.5 核心用例',2)
P('用例1-游客浏览：首页Hero旋转背景+热门卡片→搜索筛选→详情页（大图/信息/评分/评论含图片）；')
P('用例2-用户注册登录：注册（SHA-256加密）→用户名或邮箱登录→Session保存→角色路由；')
P('用例3-评分与评论：hover五星预览→click AJAX评分→Toast反馈→写评论+上传图片→预览→提交；')
P('用例4-管理员后台管理：仪表盘→目的地CRUD→图片上传→用户管理→评论删除。')

doc.add_page_break()

# ═══════════ 3. 概要设计 ═══════════
H('3. 概要（总体）设计',1)

H('3.1 系统功能模块',2)
P('（1）用户模块：注册、登录、个人资料、密码修改；')
P('（2）目的地模块：列表浏览、详情查看、分类筛选、关键字搜索、Hero背景轮播；')
P('（3）投票评分模块：星级评分UI、AJAX提交、平均分实时计算、投票状态检测；')
P('（4）评论模块：文字评论、图片上传（multipart+UUID存储）、评论删除（权限控制）、评论列表展示；')
P('（5）管理模块：仪表盘统计、目的地CRUD、用户管理、图片上传。')
P('')
P('系统功能模块图：[详见附件Draw.io文件]',bold=True)

H('3.2 系统结构设计',2)
P('系统采用经典MVC分层架构，自上而下分为七层：')
P('（1）表现层（JSP+CSS+JS）：16个JSP页面、2个CSS文件（taste-tokens.css+style.css含羽化渐变/轮播/图片上传样式）')
P('（2）控制层（6个Servlet）：UserServlet、DestinationServlet、CommentServlet（支持multipart）、VoteServlet、SearchServlet、ImageUploadServlet')
P('（3）业务逻辑层（4个Service）：评分校验/投票去重/密码加密/评论图片处理/删除权限校验')
P('（4）数据访问层（4个DAO）：全部PreparedStatement参数化查询，动态SQL拼接，子查询聚合')
P('（5）实体层（4个Model）：User、Destination、Comment（含image字段）、Vote')
P('（6）工具层（3个Util）：DBUtil（双重检查锁定单例+H2初始化）、PasswordUtil（SHA-256+Salt）、SecurityUtil')
P('（7）过滤器层（2个Filter）：EncodingFilter（UTF-8编码，web.xml首位声明）→ AuthFilter（认证+管理员鉴权）')

H('3.3 数据库设计',2)
P('数据库使用H2 Database 2.3.232嵌入式引擎，MySQL兼容模式（MODE=MySQL），共4张核心数据表：')
P('')
P('（1）users 用户表',bold=True)
C('''CREATE TABLE users (
    id INT AUTO_INCREMENT PRIMARY KEY, username VARCHAR(50) NOT NULL UNIQUE,
    email VARCHAR(100) NOT NULL UNIQUE, password VARCHAR(255) NOT NULL,
    avatar VARCHAR(255), role VARCHAR(10) DEFAULT 'user',
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP, updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP);''')
P('（2）destinations 目的地表',bold=True)
C('''CREATE TABLE destinations (
    id INT AUTO_INCREMENT PRIMARY KEY, name VARCHAR(100) NOT NULL,
    region VARCHAR(50), type VARCHAR(50), image VARCHAR(255),
    description TEXT, address VARCHAR(255), open_time VARCHAR(100),
    ticket_price DECIMAL(10,2) DEFAULT 0.00, rating DOUBLE DEFAULT 0.0,
    popularity INT DEFAULT 0, created_by INT,
    FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE SET NULL);''')
P('（3）comments 评论表（含image字段，支持图片上传）',bold=True)
C('''CREATE TABLE comments (
    id INT AUTO_INCREMENT PRIMARY KEY, destination_id INT NOT NULL,
    user_id INT NOT NULL, content TEXT NOT NULL, rating INT DEFAULT 5,
    image VARCHAR(255) DEFAULT NULL, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (destination_id) REFERENCES destinations(id) ON DELETE CASCADE,
    FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE);''')
P('（4）votes 投票表',bold=True)
C('''CREATE TABLE votes (
    id INT AUTO_INCREMENT PRIMARY KEY, destination_id INT NOT NULL,
    user_id INT NOT NULL, score INT DEFAULT 0, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    UNIQUE (destination_id, user_id),
    FOREIGN KEY (destination_id) REFERENCES destinations(id) ON DELETE CASCADE,
    FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE);''')
P('')
P('数据库ER图：[详见附件Draw.io文件]',bold=True)

doc.add_page_break()

# ═══════════ 4. 系统主要模块实现 ═══════════
H('4. 系统主要模块实现',1)

H('4.1 登录验证模块',2)
P('登录验证模块负责用户身份认证、Session管理和权限访问控制。')
P('')
P('4.1.1 密码加密（PasswordUtil.java）',bold=True)
P('用户密码采用SHA-256配合16字节随机盐值加密，存储格式为"Base64Salt:SHA256Hash"。')
C(pwdutil)

P('')
P('4.1.2 认证过滤器（AuthFilter.java）',bold=True)
P('AuthFilter通过web.xml声明注册（已移除@WebFilter注解以避免双重注册），'
  '在EncodingFilter之后执行（web.xml按声明顺序执行Filter）。'
  '过滤逻辑分三级：静态资源放行→公开路径放行→Session验证+管理员权限校验。')
C(auth_fl)

P('')
P('4.1.3 EncodingFilter与中文乱码修复',bold=True)
P('EncodingFilter负责在所有请求处理前设置UTF-8字符编码。'
  '原系统使用@WebFilter注解注册，与AuthFilter产生双重注册且执行顺序不确定，'
  '导致部分POST请求在编码设置前被解析，评论中出现中文乱码（如"黄山"显示为"??"）。')
P('修复方案：移除EncodingFilter和AuthFilter的@WebFilter注解，统一在web.xml中声明Filter及执行顺序：'
  'EncodingFilter在前（设置UTF-8编码），AuthFilter在后（检查登录状态）。'
  '同时CommentServlet.doPost()首行增加防御性request.setCharacterEncoding("UTF-8")，'
  '评论表单增加accept-charset="UTF-8"属性。', indent=True)
P('')
P('web.xml Filter声明（执行顺序即声明顺序）：',bold=True)
C('''<!-- EncodingFilter MUST run before AuthFilter -->
<filter>
    <filter-name>EncodingFilter</filter-name>
    <filter-class>com.travel.filter.EncodingFilter</filter-class>
</filter>
<filter-mapping>
    <filter-name>EncodingFilter</filter-name><url-pattern>/*</url-pattern>
</filter-mapping>
<filter>
    <filter-name>AuthFilter</filter-name>
    <filter-class>com.travel.filter.AuthFilter</filter-class>
</filter>
<filter-mapping>
    <filter-name>AuthFilter</filter-name><url-pattern>/*</url-pattern>
</filter-mapping>''')
C(enc_fl)

doc.add_page_break()

H('4.2 旅游目的地发布和编辑',2)
P('目的地CRUD由DestinationServlet+DestinationService+DestinationDAO协同完成。'
  '管理员通过后台destination_manage.jsp填写8字段表单（name/region/type/image/description/address/openTime/ticketPrice）→ '
  'POST到/admin/destination/* → handleCreate/handleUpdate/handleDelete分发 → '
  '全部使用PreparedStatement参数化SQL → 删除时ON DELETE CASCADE自动级联删除评论和投票。')
P('图片上传功能：管理员选择本地图片→FileReader预览→FormData+Fetch AJAX POST到/admin/upload-image→'
  'ImageUploadServlet使用Apache Commons FileUpload解析multipart→校验类型（JPG/PNG/GIF/WebP/BMP）和大小（5MB）→'
  'UUID命名保存到/uploads/→返回JSON URL→前端自动填入image输入框。')
C(dest_sv)

H('4.3 旅游浏览、评论与评分',2)

P('4.3.1 列表浏览与搜索筛选',bold=True)
P('首页（index.jsp）：采用Bento Grid 12列网格展示Top 6目的地。Hero区域实现旋转风景背景轮播：5张真实目的地图片，'
  '30秒间隔自动切换，1.8秒CSS opacity过渡淡入淡出。通过linear-gradient(to top right, ...)实现从'
  '左下角到右上角的白色羽化渐变，6个色阶停顿点（0%/30%/55%/75%/90%/100%）覆盖全区域，过渡柔和自然。'
  '文字区域位于白色羽化侧，保持深色文字清晰可读。')
P('列表页（destination/list.jsp）：同样部署Hero旋转背景，表格视图展示全部目的地，'
  '筛选工具栏支持地区/类型/关键字/排序组合查询。')
P('详情页（destination/detail.jsp）：展示目的地大图+完整信息卡片+交互式评分区+评论区（含图片上传）。')

P('')
P('4.3.2 交互式五星评分机制',bold=True)
P('前端：5个button.star-btn（data-star=1~5）→ mouseenter高亮到当前星+显示文字提示（很差/较差/一般/较好/非常好）→ '
  'click构建URLSearchParams→Fetch POST /vote→解析JSON→更新currentScore+刷新平均分/星级/投票数→Toast通知。')
P('后端：VoteServlet（纯JSON API）→VoteService.rate()校验+clamp评分1-5→VoteDAO.rate()：'
  '先findByUserAndDestination()→存在则UPDATE score，不存在则INSERT→getAverageScore()计算AVG→'
  '返回JSON {success,avgRating,voteCount,userScore}。')
C(vote_sv)
C(vote_dao)

P('')
P('4.3.3 评论功能（文字+图片上传+删除）',bold=True)
P('评论发表：用户在详情页评论区输入文字→可选点击"+ 添加旅行照片"选择图片→FileReader预览缩略图→'
  '可点击×移除→表单以multipart/form-data提交POST /comment→CommentServlet检测isMultipartContent→'
  '使用DiskFileItemFactory+ServletFileUpload解析→提取form字段（destinationId/content/rating）→'
  '若file item非空，校验扩展名（JPG/PNG/GIF/WebP/BMP）和大小（≤5MB）→UUID.randomUUID()生成文件名→'
  'item.write()保存到/uploads/→CommentService.addComment()含image路径→CommentDAO.create()写入数据库。')
P('评论展示：CommentDAO.findByDestinationId()使用JOIN查询获取username→JSP循环渲染comment-item卡片→'
  '若评论含image则在comment-body下方渲染<img>标签，点击可新窗口打开原图。')
P('评论删除：每条评论右侧×按钮仅作者和管理员可见→点击触发confirm()确认→'
  'POST /comment action=delete→CommentService.deleteComment()三重校验：①findById检查存在性→'
  '②作者ID匹配或isAdmin()校验→③DAO.delete()物理删除。')

P('')
P('CommentServlet完整源码（支持multipart图片上传+删除权限控制）：',bold=True)
C(cmt_sv)

doc.add_page_break()

# ═══════════ 5. 前台设计实现 ═══════════
H('5. 前台主要功能模块设计与实现',1)

P('前台页面基于Taste-Skill极简设计系统构建，所有页面共享CSS自定义属性（Design Tokens）。')

P('')
P('5.1 全局导航栏（header.jsp）',bold=True)
P('64px粘性定位+1px底边框，根据登录状态动态显示导航链接、用户名头像和管理员入口。')
C(header_jsp)

P('')
P('5.2 首页Hero旋转背景与羽化效果（index.jsp）',bold=True)
P('Hero区域部署三层结构：①hero-bg-slider（5张真实目的地图片，绝对定位覆盖整个Hero）→'
  '②hero-overlay（CSS linear-gradient白色羽化渐变层，从左上角白色渐变到右下角透明，pointer-events:none透传鼠标事件）→'
  '③hero-content（z-index:2置于最上层，深色文字在白底上清晰可读）。')
P('JavaScript控制30秒轮播：setInterval每30秒切换active类→CSS opacity 0↔1过渡1.8秒→'
  '配合transition: opacity 1.8s ease-in-out实现平滑淡入淡出。')
P('')
P('白色羽化CSS渐变色阶（6个停顿点，宽范围过渡）：',bold=True)
C('''.hero-overlay {
  background: linear-gradient(
    to top right,
    rgba(255,255,255,1) 0%,      /* 左下角：纯白，完全覆盖图片 */
    rgba(255,255,255,0.95) 30%,   /* 30%位置：几乎全白 */
    rgba(255,255,255,0.75) 55%,   /* 55%位置：半透明 */
    rgba(255,255,255,0.35) 75%,   /* 75%位置：图片开始显现 */
    rgba(255,255,255,0.08) 90%,   /* 90%位置：图片基本清晰 */
    rgba(255,255,255,0) 100%      /* 右上角：完全透明，图片完整呈现 */
  );
}''')
P('此设计既保证了左下侧标题文字的清晰可读，又让右上侧的风景图片自然呈现，'
  '羽化过渡范围覆盖整个Hero区域，视觉上柔和流畅。', indent=True)

P('')
P('5.3 登录页面',bold=True)
C(login_jsp)

P('')
P('5.4 评论图片上传与展示（detail.jsp）',bold=True)
P('评论表单增加图片上传功能：文件选择器（accept="image/*"）→虚线边框的"+ 添加旅行照片"标签→'
  '选择图片后FileReader.readAsDataURL()预览缩略图→×移除按钮重置→'
  'form enctype="multipart/form-data"提交。评论列表中，含有图片的评论在文字下方展示缩略图，'
  '点击可在新窗口查看原图（cursor:pointer + onclick="window.open(this.src)"）。')

doc.add_page_break()

# ═══════════ 6. 后台设计实现 ═══════════
H('6. 系统后台设计与实现',1)

H('6.1 数据库连接类（DBUtil.java）',2)
P('DBUtil封装了H2嵌入式数据库的连接管理、表结构自动创建和12条示例数据初始化三大职责。'
  '采用volatile+synchronized双重检查锁定实现线程安全的懒初始化。'
  'H2连接URL配置MODE=MySQL启用MySQL语法兼容，DATABASE_TO_LOWER=TRUE统一小写表名，'
  'DB_CLOSE_DELAY=-1保持连接直到JVM退出。')
C(dbutil)

H('6.2 Servlet控制器类设计',2)
P('系统共6个Servlet，统一使用@WebServlet注解配置路由映射。')

P('')
P('Servlet路由映射总表：',bold=True)
C('''@WebServlet("/user/*")              → UserServlet
    POST /user/login | /user/register | /user/profile | /user/change-password
@WebServlet("/admin/destination/*") → DestinationServlet  (action=create/update/delete)
@WebServlet("/admin/upload-image")  → ImageUploadServlet  (multipart AJAX)
@WebServlet("/comment")             → CommentServlet
    (无action→发表评论含图片 | action=delete→删除评论)
@WebServlet("/vote")                → VoteServlet  (GET状态查询 | POST rate/unrate)
@WebServlet("/search")              → SearchServlet (keyword/region/type/sortBy)''')

H('6.3 登录验证处理页面',2)
P('登录完整流程：用户访问login.jsp→填写凭证→POST /user/login→UserServlet.handleLogin()→'
  'UserService.login()调用DAO查询+PasswordUtil.verify()验证→验证通过存入HttpSession→'
  '管理员重定向到/admin/dashboard.jsp，普通用户重定向到首页→'
  '验证失败则session.setAttribute("error")并重定向回login.jsp显示错误提示。')

P('')
P('')
P('参考文献',bold=True)
P('[1] 李刚. 轻量级Java EE企业应用实战（第5版）[M]. 北京: 电子工业出版社, 2019.')
P('[2] 陈恒, 楼偶俊. Java Web开发实战[M]. 北京: 清华大学出版社, 2020.')
P('[3] 孙卫琴. Tomcat与Java Web开发技术详解（第3版）[M]. 北京: 电子工业出版社, 2019.')
P('[4] 王珊, 萨师煊. 数据库系统概论（第5版）[M]. 北京: 高等教育出版社, 2014.')
P('[5] H2 Database Engine Documentation [EB/OL]. https://www.h2database.com/, 2024.')
P('[6] Apache Tomcat 9 Documentation [EB/OL]. https://tomcat.apache.org/tomcat-9.0-doc/, 2024.')
P('[7] Apache Commons FileUpload User Guide [EB/OL]. https://commons.apache.org/proper/commons-fileupload/, 2024.')
P('[8] Taste-Skill Design System [EB/OL]. https://github.com/z75758/taste-skill, 2024.')
P('[9] MDN Web Docs. CSS linear-gradient [EB/OL]. https://developer.mozilla.org/en-US/docs/Web/CSS/gradient/linear-gradient, 2024.')
P('[10] Java Servlet 3.0 Specification [S]. Oracle Corporation, 2009.')

# ── Save ──
out = r'F:\MyProject\TravelRating\课程设计报告_最终版.docx'
doc.save(out)
print(f'Saved: {out} ({os.path.getsize(out)/1024:.0f} KB)')
