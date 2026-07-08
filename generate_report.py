"""
Generate the Fourth Section course design report for TravelRating system.
Template: Web开发技术课程设计报告 第四部分
Output: report_section4.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ── Helper functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    return h

def add_para(text, bold=False, indent=True, font_size=12, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if align is not None:
        p.alignment = align
    return p

def add_code(code_text):
    """Add a code block with monospace font."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(51, 51, 51)
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(2)
    gap.paragraph_format.space_after = Pt(2)

def add_bullet(text, indent_cm=0.74):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── Read source code files ──
base = r'F:\MyProject\TravelRating\src\main'

def read_src(path):
    full = os.path.join(base, path)
    if os.path.exists(full):
        with open(full, encoding='utf-8') as f:
            return f.read()
    return ''

# Java source
dbutil_code = read_src('java/com/travel/util/DBUtil.java')
password_util = read_src('java/com/travel/util/PasswordUtil.java')
user_servlet = read_src('java/com/travel/controller/UserServlet.java')
dest_servlet = read_src('java/com/travel/controller/DestinationServlet.java')
vote_servlet = read_src('java/com/travel/controller/VoteServlet.java')
comment_servlet = read_src('java/com/travel/controller/CommentServlet.java')
search_servlet = read_src('java/com/travel/controller/SearchServlet.java')
image_upload = read_src('java/com/travel/controller/ImageUploadServlet.java')
auth_filter = read_src('java/com/travel/filter/AuthFilter.java')
user_dao = read_src('java/com/travel/dao/UserDAO.java')
dest_dao = read_src('java/com/travel/dao/DestinationDAO.java')
vote_dao = read_src('java/com/travel/dao/VoteDAO.java')
vote_service = read_src('java/com/travel/service/VoteService.java')

# JSP pages
index_jsp = read_src('webapp/index.jsp')
login_jsp = read_src('webapp/login.jsp')
header_jsp = read_src('webapp/header.jsp')
detail_jsp = read_src('webapp/destination/detail.jsp')
list_jsp = read_src('webapp/destination/list.jsp')
dest_mgmt = read_src('webapp/admin/destination_manage.jsp')
dashboard_jsp = read_src('webapp/admin/dashboard.jsp')

# Config
webxml = read_src('webapp/WEB-INF/web.xml')

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('第四部分：课程设计内容（步骤及程序）')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('项目名称：旅游地点评选与互动系统设计与实现')
run.font.size = Pt(14)

add_para('')
add_para('摘要：本项目为旅游地点评选与互动系统，是一个基于JSP + Servlet + JDBC技术栈的'
         'B/S架构Web应用程序。系统面向游客与管理员两类用户，提供旅游目的地浏览、'
         '多条件搜索筛选、交互式五星评分、评论互动、后台管理等完整功能。'
         '技术实现上采用MVC分层架构（Model-View-Controller），以H2嵌入式数据库实现零安装部署，'
         '采用SHA-256+随机盐值加密保障用户密码安全，前端基于Taste-Skill极简设计系统'
         '实现暗色模式和响应式布局。系统代码约3500行，包含6个Servlet控制器、4个DAO数据访问类、'
         '4个Service业务逻辑类、4个实体模型类、2个Filter过滤器、3个工具类及16个JSP页面。', indent=True)
add_para('')
add_para('关键词：旅游目的地评选；JSP；Servlet；JDBC；H2数据库；MVC架构；Taste-Skill设计系统', bold=True)

doc.add_page_break()

# ============================================================
# 1. 系统概述
# ============================================================
add_heading_styled('1. 系统概述', level=1)

add_heading_styled('1.1 设计目的', level=2)
add_para('随着我国旅游业的蓬勃发展，游客对于旅游目的地的信息获取和评价反馈需求日益增长。'
         '传统的旅游信息网站多采用静态内容展示方式，缺乏用户互动和评价机制，'
         '难以满足现代游客对个性化、社交化旅游信息服务的需求。'
         '本课程设计旨在综合运用Web开发技术，完成一个集目的地展示、搜索筛选、'
         '星级评分、评论互动于一体的完整Web应用系统。具体设计目的包括：')
add_para('（1）综合运用JSP、Servlet、JDBC等Java Web核心技术，构建一个功能完备的B/S架构应用系统；')
add_para('（2）深入实践MVC分层架构设计思想，将数据访问层（DAO）、业务逻辑层（Service）、'
         '控制层（Servlet）和表现层（JSP）进行清晰的职责分离；')
add_para('（3）掌握关系型数据库的设计与实现，包括ER图设计、表结构规范化、'
         'SQL查询优化以及外键约束与级联操作的应用；')
add_para('（4）实现完整的用户认证与权限管理体系，深入理解Session机制、Filter过滤器链'
         '以及SHA-256+Salt密码加密策略；')
add_para('（5）应用Taste-Skill极简设计系统，实践CSS自定义属性、响应式布局、'
         '暗色模式适配等现代前端设计理念；')
add_para('（6）掌握AJAX异步交互技术，实现无刷新星级评分、异步图片上传等功能，提升用户体验。')

add_heading_styled('1.2 项目概况', level=2)
add_para('旅游地点评选与互动系统（TravelRating）是一个基于Java Web技术栈的在线旅游信息服务平台，'
         '系统面向两类用户角色：普通用户（游客）和系统管理员。', indent=True)

add_para('普通用户功能：浏览旅游目的地列表、按地区/类型/评分/票价进行多条件筛选、'
         '关键字搜索目的地、查看目的地详情信息（大图、文字描述、地址、开放时间、门票价格、'
         '综合评分、用户评价）、通过交互式五星组件对目的地进行1-5分评分、'
         '发表文字评论分享旅行体验、管理个人资料（编辑用户名/邮箱/头像、修改密码）。')

add_para('管理员功能：后台仪表盘查看系统统计数据摘要（目的地总数、用户总数、总投票数）、'
         '目的地信息的完整CRUD管理（创建/编辑/删除）、用户列表查看与管理、'
         '目的地封面图片的AJAX异步本地上传。')

add_para('技术架构：系统严格遵循MVC三层架构。表现层采用JSP页面 + Taste-Skill CSS设计系统'
         ' + Vanilla JavaScript实现页面渲染与用户交互；控制层由6个Servlet类组成，'
         '通过@WebServlet注解配置URL映射，接收HTTP请求并调用Service层处理业务逻辑；'
         '业务逻辑层包含4个Service类，封装投票去重、评分计算、密码验证等核心业务规则；'
         '数据访问层包含4个DAO类，使用JDBC PreparedStatement执行参数化SQL查询，'
         '从根源上防止SQL注入攻击。数据库采用H2 Embedded嵌入式数据库，'
         '以MySQL兼容模式运行（MODE=MySQL），数据文件存储于项目./data/目录，'
         '实现免安装、开箱即用的部署体验。前端采用Taste-Skill极简设计系统，'
         '以暖白色背景、1px细边框、8px统一圆角、微妙阴影效果和流畅过渡动画，'
         '营造现代简约的视觉风格。', indent=True)

doc.add_page_break()

# ============================================================
# 2. 系统分析
# ============================================================
add_heading_styled('2. 系统分析', level=1)

add_heading_styled('2.1 引言', level=2)
add_para('本节对旅游地点评选与互动系统进行全面的系统分析，'
         '从任务概述、功能/非功能需求、运行环境规定和使用场景用例四个维度，'
         '明确系统需要"做什么"以及"做到什么程度"，'
         '为后续的系统概要设计和详细编码实现奠定基础和提供方向指引。')

add_heading_styled('2.2 任务概述', level=2)
add_para('本系统的核心开发任务包括：')
add_para('（1）搭建基于Maven的Java Web项目骨架，配置Maven Compiler Plugin和'
         'tomcat7-maven-plugin嵌入式Tomcat运行环境，JDK版本设为11；')
add_para('（2）设计并实现数据库表结构，包含users（用户表）、destinations（目的地表）、'
         'comments（评论表）、votes（投票表）共4张数据表，定义完整的主键、外键、'
         'UNIQUE约束和级联删除策略；')
add_para('（3）实现用户注册与登录功能，采用SHA-256 + 16字节随机盐值（Base64编码）'
         '加密存储密码，通过AuthFilter实现基于Session的登录认证拦截；')
add_para('（4）实现旅游目的地的增删改查（CRUD）管理功能，管理员可通过后台页面'
         '进行目的地的创建、编辑和删除操作；')
add_para('（5）实现交互式五星评分机制，用户通过点击星星进行1-5分评分，'
         '每用户每目的地仅可评分一次（UNIQUE约束 + 应用层UPDATE逻辑），'
         '系统实时计算并更新目的地平均评分；')
add_para('（6）实现评论功能，登录用户可在目的地详情页发表文字评论，'
         '评论列表按时间倒序展示，包含评论者用户名和时间戳；')
add_para('（7）实现多条件搜索与筛选功能，支持按关键字模糊匹配名称和描述、'
         '按地区筛选、按类型筛选、按评分排序、按热度排序、按最新排序；')
add_para('（8）实现图片本地上传功能，通过FormData + Fetch API进行AJAX异步上传，'
         '服务端校验文件类型和大小，生成UUID文件名保存至/uploads/目录；')
add_para('（9）基于Taste-Skill设计系统完成全部前端页面样式设计，实现统一的视觉风格；')
add_para('（10）实现响应式布局（768px断点）和暗色模式（prefers-color-scheme）自适应。')

add_heading_styled('2.3 需求规定', level=2)

add_para('一、功能需求', bold=True)
add_bullet('用户模块：用户注册（校验用户名和邮箱唯一性）、登录（支持用户名或邮箱登录）、'
           '个人资料编辑（用户名/邮箱/头像URL）、密码修改（需验证原密码）、安全退出（Session销毁）；')
add_bullet('目的地模块：首页Bento Grid卡片式展示热门目的地、列表页表格视图展示全部目的地、'
           '按地区（华东/华北/华南/西南/西北/华中/东北）下拉筛选、按类型（自然风光/历史古迹/'
           '海岛度假/美食之旅/城市观光）下拉筛选、关键字模糊搜索、按评分/热度/最新排序、'
           '目的地详情页（大幅封面图、文字描述、地址、开放时间、门票价格、综合评分、访问热度）；')
add_bullet('投票评分模块：交互式五星评分组件（hover高亮动画、click提交AJAX请求）、'
           '每用户每目的地仅可评分一次（已评分则UPDATE更新、未评分则INSERT新增）、'
           '可取消评分（DELETE删除记录）、评分后实时更新平均分和星级显示、Toast消息反馈；')
add_bullet('评论模块：登录用户发表文字评论、评论列表按时间倒序展示（含用户名头像首字母和时间）、'
           '评论与评分功能相互独立（可只评分不评论、也可只评论不评分）；')
add_bullet('管理模块：后台仪表盘统计卡片（目的地总数/用户总数/总投票数）、'
           '目的地CRUD管理（表单创建、编辑、删除确认）、用户列表管理、'
           '目的地封面图片本地上传（JPG/PNG/GIF/WebP/BMP，最大5MB）；')
add_bullet('搜索模块：首页全局搜索栏（关键字+地区+类型）、列表页筛选工具栏、搜索结果页面展示。')

add_para('')
add_para('二、非功能需求', bold=True)
add_bullet('安全性：密码采用SHA-256 + 16字节随机盐值加密存储（格式："Base64Salt:SHA256Hash"），'
           '全部数据库操作使用PreparedStatement参数化查询防止SQL注入，'
           'Session超时30分钟后自动失效，AuthFilter拦截所有非公开路径的未登录请求；')
add_bullet('可用性：界面布局简洁直观，操作具有即时反馈（成功/失败Toast消息提示），'
           '重要操作（删除）提供二次确认对话框，表单提交提供前端校验（密码一致性、必填项检查）；')
add_bullet('兼容性：支持Chrome/Firefox/Edge/Safari等主流现代浏览器，'
           '768px断点以下自适应为单列移动端布局，触控设备友好（按钮尺寸≥44px）；')
add_bullet('可维护性：严格遵循MVC分层架构（代码按model/dao/service/controller/filter/util分包），'
           '类和方法命名语义化，关键逻辑添加中文注释说明，DAO层统一使用try-with-resources管理JDBC资源。')

add_heading_styled('2.4 运行环境规定', level=2)
add_para('硬件环境：', bold=True)
add_bullet('CPU：Intel Core i5 或同等性能处理器及以上；')
add_bullet('内存：8GB 及以上（JVM堆内存建议≥512MB）；')
add_bullet('硬盘：500MB 可用空间（含JDK、Maven依赖及项目文件）。')

add_para('')
add_para('软件环境：', bold=True)
add_bullet('操作系统：Windows 10/11 64位；')
add_bullet('JDK版本：OpenJDK 11（编译源码级别 source/target = 11）；')
add_bullet('构建工具：Apache Maven 3.9+，使用tomcat7-maven-plugin 2.2嵌入式运行；')
add_bullet('应用服务器：Apache Tomcat 9.0.97（通过Maven插件内嵌运行，端口8080）；')
add_bullet('数据库：H2 Database Engine 2.3.232，MySQL兼容模式（MODE=MySQL），'
           '文件存储在./data/travel_db.mv.db；')
add_bullet('开发工具：VS Code / IntelliJ IDEA；')
add_bullet('浏览器：Google Chrome 120+（推荐），Firefox 115+，Edge 120+。')

add_heading_styled('2.5 需求分析（核心用例）', level=2)
add_para('通过对旅游信息服务的典型应用场景进行梳理，提炼出以下四个核心用例：')

add_para('用例1：游客浏览与搜索目的地', bold=True)
add_para('前置条件：用户打开网站首页。'
         '基本流程：① 首页展示热门目的地卡片（Bento Grid布局）和系统统计数据 → '
         '② 用户可在搜索栏输入关键字、选择地区/类型进行筛选 → '
         '③ 点击"浏览全部"进入列表页，查看完整目的地表格 → '
         '④ 点击任意目的地卡片进入详情页 → '
         '⑤ 详情页展示大幅封面图、完整描述、地址、开放时间、票价、综合评分（星级+数值）、'
         '投票人数和用户评论列表。')

add_para('用例2：用户注册与登录', bold=True)
add_para('前置条件：用户点击"Sign In"或"Join Us"按钮。'
         '基本流程：① 注册：填写用户名/邮箱/密码/确认密码 → 前端校验两次密码一致性 → '
         '提交后系统校验用户名和邮箱唯一性 → 密码经SHA-256+Salt加密后存入数据库 → '
         '注册成功跳转登录页 → ② 登录：输入用户名或邮箱及密码 → 系统查询用户并验证密码哈希值 → '
         '验证通过将User对象存入HttpSession → 管理员自动跳转后台仪表盘，普通用户跳转首页。')

add_para('用例3：星级评分与评论互动', bold=True)
add_para('前置条件：用户已登录并进入目的地详情页。'
         '基本流程：① 用户将鼠标悬停在五星评分组件上，星星高亮变色为金黄色(#e6a817)并显示'
         '评分文字提示（很差/较差/一般/较好/非常好） → '
         '② 点击目标星级，JavaScript触发AJAX POST请求到/vote → '
         '③ VoteServlet接收参数，VoteService检查当前用户是否已对该目的地评过分 → '
         '④ 若未评分：INSERT新vote记录；若已评分：UPDATE已有记录的score值 → '
         '⑤ VoteService调用DestinationDAO.updateRating()重新计算平均分 → '
         '⑥ 返回JSON（含新平均分和投票数），前端实时更新星级显示和统计数字 → '
         '⑦ 用户在评论区输入文字内容，点击提交 → CommentServlet处理 → '
         '页面刷新后评论列表显示最新评论（按时间倒序）。')

add_para('用例4：管理员后台管理', bold=True)
add_para('前置条件：管理员账号（admin）登录系统。'
         '基本流程：① 登录后自动跳转后台仪表盘 → 查看目的地总数、用户总数、总投票数 → '
         '② 进入目的地管理页 → 点击"添加目的地"填写表单（名称/地区/类型/图片/描述/地址/'
         '开放时间/票价） → 可点击"Choose File"上传本地图片 → 提交表单 → '
         '③ DestinationServlet接收POST创建请求，验证数据完整性和管理员权限 → '
         '调用DestinationDAO.create()写入数据库 → 重定向回管理页显示新记录 → '
         '④ 编辑：点击编辑按钮修改已有目的地信息 → 提交更新 → '
         '⑤ 删除：点击删除按钮，确认后执行DELETE操作（级联删除关联评论和投票）。')

doc.add_page_break()

# ============================================================
# 3. 概要（总体）设计
# ============================================================
add_heading_styled('3. 概要（总体）设计', level=1)

add_heading_styled('3.1 系统功能模块设计', level=2)
add_para('根据需求分析结果，旅游地点评选与互动系统划分为五大功能模块：')
add_para('')
add_para('（1）用户模块：提供注册、登录、个人资料管理、密码修改、安全退出功能。'
         '是系统的基础模块，为其他功能提供身份认证和权限控制支持。')
add_para('（2）目的地模块：提供目的地列表浏览、详情查看、多条件筛选（地区/类型）、'
         '关键字模糊搜索、多维度排序（热度/评分/最新）。是系统最核心的信息展示模块。')
add_para('（3）投票评分模块：提供交互式五星评分UI、AJAX异步评分提交/修改/取消、'
         '评分状态检测（该用户是否已评分）、平均评分实时计算与更新。是系统最具特色的互动模块。')
add_para('（4）评论模块：提供评论发表表单、评论列表展示（含用户信息）、'
         '评论时间排序。评论与评分模块相互独立，用户可灵活选择互动方式。')
add_para('（5）管理模块：提供后台数据统计仪表盘、目的地信息完整CRUD管理、'
         '用户信息管理、封面图片AJAX异步上传。是系统运维管理的核心入口。')

add_para('')
add_para('系统功能模块图如下：', indent=True)
add_para('[图1-2 系统功能模块图 —— 详见附件Draw.io文件]', bold=True, indent=False)

add_heading_styled('3.2 系统结构设计', level=2)
add_para('系统采用经典的MVC（Model-View-Controller）分层架构，自上而下分为七个层次：')

add_para('')
add_para('（1）表现层（View — JSP + CSS + JavaScript）', bold=True)
add_para('由16个JSP页面、2个CSS样式文件（taste-tokens.css + style.css）和页面内嵌JavaScript脚本组成。'
         '负责页面布局渲染、用户交互响应和前端数据校验。'
         '通过JSP脚本片段（Scriptlet）和EL表达式动态展示后端传递的数据。')

add_para('（2）控制层（Controller — Servlet）', bold=True)
add_para('由6个Servlet类组成，全部使用@WebServlet注解声明式配置URL路由映射，'
         '无需在web.xml中手动配置<servlet>和<servlet-mapping>元素。'
         '接收HTTP请求参数，调用Service层执行业务逻辑，根据处理结果决定重定向（PRG模式）'
         '或返回JSON数据（AJAX请求）。')

add_para('（3）业务逻辑层（Service）', bold=True)
add_para('由4个Service类（UserService、DestinationService、VoteService、CommentService）组成。'
         '封装核心业务规则与流程：投票去重检查（INSERT or UPDATE逻辑）、'
         '评分范围校验（clamp 1-5）、密码加密验证（SHA-256+Salt）、'
         '目的地平均评分同步更新（评分变更后自动触发refreshRating）。'
         '作为Controller和DAO之间的业务抽象层。')

add_para('（4）数据访问层（DAO）', bold=True)
add_para('由4个DAO类（UserDAO、DestinationDAO、CommentDAO、VoteDAO）组成。'
         '封装全部数据库CRUD操作，所有SQL查询均使用JDBC PreparedStatement进行参数化，'
         '从根源上杜绝SQL注入攻击。资源管理统一采用try-with-resources自动关闭机制。'
         '支持动态SQL拼接（搜索筛选的多条件WHERE子句）。')

add_para('（5）实体层（Model/POJO）', bold=True)
add_para('由4个实体类（User、Destination、Comment、Vote）组成。'
         '纯POJO（Plain Old Java Object），字段与数据库表列一一对应。'
         '通过Getter/Setter方法访问属性，部分实体类包含便捷方法（如Destination.getTicketPriceDisplay()、'
         'User.getInitial()、User.isAdmin()）。')

add_para('（6）工具层（Util）', bold=True)
add_para('包含3个工具类：DBUtil（数据库连接池管理，双重检查锁定单例初始化，'
         '首次调用时自动建表并插入12条中国著名旅游目的地示例数据）、'
         'PasswordUtil（SHA-256 + 随机盐值密码哈希与验证）、SecurityUtil（输入安全过滤）。')

add_para('（7）过滤器层（Filter）', bold=True)
add_para('包含2个过滤器：AuthFilter（基于Session的登录认证拦截，区分公开路径与受保护路径，'
         '管理员路径额外权限校验）、EncodingFilter（统一设置请求和响应的UTF-8字符编码）。')

add_para('')
add_para('项目完整目录结构（共约50个源文件）：', bold=True)
add_code('''TravelRating/
├── pom.xml                             # Maven项目配置，JDK11+Tomcat9+H2
├── .gitignore
├── data/                               # H2数据库文件（自动生成）
└── src/main/
    ├── java/com/travel/
    │   ├── controller/                 # 控制层（6个Servlet）
    │   │   ├── UserServlet.java           用户注册/登录/登出/资料/改密
    │   │   ├── DestinationServlet.java    目的地CRUD管理（管理员）
    │   │   ├── CommentServlet.java        评论发表与删除
    │   │   ├── VoteServlet.java           星级评分（AJAX JSON接口）
    │   │   ├── SearchServlet.java         搜索筛选（GET+POST）
    │   │   └── ImageUploadServlet.java    图片本地上传（AJAX）
    │   ├── dao/                        # 数据访问层（4个DAO）
    │   │   ├── UserDAO.java              用户表CRUD + 按用户名/邮箱查询
    │   │   ├── DestinationDAO.java       目的地表CRUD + 多维搜索 + 评分刷新
    │   │   ├── CommentDAO.java           评论表CRUD + 按目的地查询+用户名JOIN
    │   │   └── VoteDAO.java              评分UPSERT + 取消评分 + AVG计算
    │   ├── service/                    # 业务逻辑层（4个Service）
    │   │   ├── UserService.java          注册唯一性检查、登录验证、改密
    │   │   ├── DestinationService.java   目的地CRUD + 热度增加 + 评分刷新
    │   │   ├── CommentService.java       评论增删 + 评分校验
    │   │   └── VoteService.java          评分校验 + DAO调用 + 评分同步
    │   ├── model/                      # 实体层（4个POJO）
    │   │   ├── User.java                 id/username/email/password/role/...
    │   │   ├── Destination.java          id/name/region/type/rating/...
    │   │   ├── Comment.java              id/destinationId/userId/content/...
    │   │   └── Vote.java                 id/destinationId/userId/score
    │   ├── filter/                     # 过滤器层（2个Filter）
    │   │   ├── AuthFilter.java           登录认证 + 管理员权限拦截
    │   │   └── EncodingFilter.java       UTF-8编码设置
    │   └── util/                       # 工具层（3个Util）
    │       ├── DBUtil.java               H2连接管理 + 自动建表 + 示例数据
    │       ├── PasswordUtil.java         SHA-256 + Salt 密码哈希
    │       └── SecurityUtil.java         输入安全过滤
    └── webapp/
        ├── WEB-INF/web.xml              Servlet 3.0声明，metadata-complete=false
        ├── css/
        │   ├── taste-tokens.css         Taste-Skill设计令牌（CSS变量）
        │   └── style.css                全局样式 + 星级评分组件样式
        ├── js/                          （预留）公共JavaScript
        ├── uploads/                     图片上传目录
        ├── admin/
        │   ├── dashboard.jsp            后台仪表盘
        │   ├── destination_manage.jsp   目的地管理（CRUD + 图片上传）
        │   └── user_manage.jsp          用户管理
        ├── destination/
        │   ├── list.jsp                 目的地列表 + 搜索筛选
        │   └── detail.jsp              目的地详情 + 星级评分 + 评论
        ├── user/
        │   └── profile.jsp             个人中心（资料编辑+密码修改）
        ├── index.jsp                   首页
        ├── header.jsp                  全局导航栏
        ├── footer.jsp                  全局页脚
        ├── login.jsp                   登录页面
        └── register.jsp                注册页面''')

add_heading_styled('3.3 数据库设计', level=2)
add_para('系统使用H2 Database 2.3.232嵌入式数据库引擎，以MySQL兼容模式运行'
         '（连接URL中设置MODE=MySQL），数据库文件存储在项目./data/目录下。'
         '数据库名为travel_db，共设计4张核心数据表：users（用户表）、'
         'destinations（目的地表）、comments（评论表）、votes（投票/评分表）。'
         '各表之间通过外键约束保证数据参照完整性，通过ON DELETE CASCADE/SET NULL'
         '实现级联删除/置空策略。')

add_para('')
add_para('H2数据库连接配置（摘自DBUtil.java）：', bold=True)
add_code('''// H2 Embedded Database 连接配置
private static final String DRIVER = "org.h2.Driver";
private static final String URL = "jdbc:h2:file:./data/travel_db"
        + ";MODE=MySQL"                // 启用MySQL SQL语法兼容
        + ";DATABASE_TO_LOWER=TRUE"    // 表名和列名自动转为小写
        + ";DB_CLOSE_DELAY=-1";        // JVM退出前保持数据库连接
private static final String USER = "sa";
private static final String PASSWORD = "";''')

add_para('')
add_para('（1）users 用户表 —— 存储所有注册用户的基本信息和认证凭证', bold=True)
add_code('''CREATE TABLE users (
    id          INT AUTO_INCREMENT PRIMARY KEY,    -- 用户ID，自增主键
    username    VARCHAR(50)  NOT NULL UNIQUE,      -- 用户名，唯一索引，用于登录和显示
    email       VARCHAR(100) NOT NULL UNIQUE,      -- 邮箱，唯一索引，支持邮箱登录
    password    VARCHAR(255) NOT NULL,             -- 密码哈希：Base64Salt:SHA256Hash
    avatar      VARCHAR(255) DEFAULT NULL,         -- 用户头像URL（可选）
    role        VARCHAR(10)  DEFAULT 'user',       -- 角色：user（普通）/ admin（管理员）
    created_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP,  -- 注册时间
    updated_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP   -- 最后更新时间
);''')

add_para('（2）destinations 目的地表 —— 核心业务表，存储旅游目的地的完整信息', bold=True)
add_code('''CREATE TABLE destinations (
    id           INT AUTO_INCREMENT PRIMARY KEY,  -- 目的地ID，自增主键
    name         VARCHAR(100) NOT NULL,           -- 目的地名称
    region       VARCHAR(50)  NOT NULL,           -- 所属地区：华东/华北/华南/西南/西北/华中/东北
    type         VARCHAR(50)  NOT NULL,           -- 景点类型：自然风光/历史古迹/海岛度假/美食之旅/城市观光
    image        VARCHAR(255) DEFAULT NULL,       -- 封面图片URL（支持外部链接或本地/uploads/路径）
    description  TEXT,                             -- 详细文字描述
    address      VARCHAR(255) DEFAULT NULL,       -- 详细地址
    open_time    VARCHAR(100) DEFAULT NULL,       -- 开放时间（如"06:00-17:00"或"全天开放"）
    ticket_price DECIMAL(10,2) DEFAULT 0.00,     -- 门票价格（0.00表示免费）
    rating       DOUBLE DEFAULT 0.0,              -- 综合评分（派生属性，由AVG(votes.score)计算得到）
    popularity   INT DEFAULT 0,                   -- 热度计数（每次查看详情自动+1）
    created_by   INT DEFAULT NULL,                -- 创建者ID（外键→users.id，管理员发布）
    created_at   TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at   TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE SET NULL
);''')

add_para('（3）comments 评论表 —— 存储用户对目的地的文字评价', bold=True)
add_code('''CREATE TABLE comments (
    id             INT AUTO_INCREMENT PRIMARY KEY,  -- 评论ID
    destination_id INT NOT NULL,                     -- 关联目的地（外键）
    user_id        INT NOT NULL,                     -- 评论者ID（外键）
    content        TEXT NOT NULL,                    -- 评论正文（TEXT类型支持长文本）
    rating         INT DEFAULT 5,                    -- 评论附带的评分（1-5），默认5分
    created_at     TIMESTAMP DEFAULT CURRENT_TIMESTAMP,  -- 评论时间
    FOREIGN KEY (destination_id) REFERENCES destinations(id) ON DELETE CASCADE,
    FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
);''')

add_para('（4）votes 投票/评分表 —— 存储用户对目的地的星级评分，核心约束保证一用户一票', bold=True)
add_code('''CREATE TABLE votes (
    id             INT AUTO_INCREMENT PRIMARY KEY,   -- 投票记录ID
    destination_id INT NOT NULL,                      -- 目的地ID（外键）
    user_id        INT NOT NULL,                      -- 投票用户ID（外键）
    score          INT DEFAULT 0,                     -- 评分值（1-5）
    created_at     TIMESTAMP DEFAULT CURRENT_TIMESTAMP, -- 投票时间
    UNIQUE (destination_id, user_id),                 -- ★核心约束：每用户每目的地仅一条记录
    FOREIGN KEY (destination_id) REFERENCES destinations(id) ON DELETE CASCADE,
    FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
);''')

add_para('')
add_para('数据库初始化时自动插入12条中国著名旅游目的地示例数据'
         '（含黄山、故宫、三亚亚龙湾、丽江古城、九寨沟、西湖、兵马俑、成都锦里、'
         '张家界、布达拉宫、呼伦贝尔大草原、鼓浪屿）和1个管理员账号'
         '（用户名admin，密码admin123的SHA-256+Salt哈希值）。', indent=True)
add_para('')
add_para('数据库ER图：', indent=True)
add_para('[图2-3 局部ER图 / 图2-4 全局ER图 —— 详见附件Draw.io文件]', bold=True, indent=False)

doc.add_page_break()

# ============================================================
# 4. 系统主要模块实现
# ============================================================
add_heading_styled('4. 系统主要模块实现', level=1)

add_heading_styled('4.1 登录验证模块', level=2)
add_para('登录验证模块负责用户身份认证、Session会话管理和权限访问控制，是系统安全体系的核心。')

add_para('')
add_para('4.1.1 密码加密与验证（PasswordUtil.java）', bold=True)
add_para('用户密码绝不使用明文存储。系统采用SHA-256算法配合16字节随机盐值进行加密：')
add_para('（1）注册/修改密码时：使用SecureRandom生成16字节随机盐 → Base64编码为字符串 → '
         '将盐值字符串与明文密码拼接 → 对整个字符串计算SHA-256哈希 → '
         '最终存储格式为"Base64Salt:SHA256Hash"（以冒号分隔）；')
add_para('（2）登录验证时：从数据库取出存储的"salt:hash"字符串 → 按冒号分割提取盐值和原哈希 → '
         '将提取的盐值+用户输入的密码拼接后计算SHA-256 → 与数据库中的哈希值比对，'
         '一致则验证通过。')

add_para('')
add_para('PasswordUtil工具类核心代码：', bold=True)
add_code(password_util)

add_para('')
add_para('4.1.2 用户认证处理（UserServlet.java）', bold=True)
add_para('UserServlet负责处理用户注册、登录、登出、个人资料编辑和密码修改共5类请求。'
         '通过@WebServlet("/user/*")注解映射路径，在doPost()方法中根据PathInfo分派到具体处理方法。')

add_para('实现要点：', bold=True)
add_para('（1）handleLogin()：接收login（用户名或邮箱）和password参数 → '
         '调用UserService.login()进行身份验证 → 验证成功则将User对象存入HttpSession → '
         '根据用户角色重定向：管理员跳转/admin/dashboard.jsp，普通用户跳转首页；')
add_para('（2）handleRegister()：接收username/email/password/confirmPassword参数 → '
         '调用UserService.register()进行用户名邮箱唯一性检查和密码一致性验证 → '
         '通过后调用PasswordUtil.hash()加密密码 → 调用UserDAO.create()写入数据库；')
add_para('（3）handleLogout()：获取当前HttpSession → 调用session.invalidate()销毁会话 → '
         '重定向至首页；')
add_para('（4）handleProfile()：验证登录状态 → 更新用户名和邮箱 → 同步更新Session中的User对象；')
add_para('（5）handleChangePassword()：验证原密码正确性 → 对新密码进行SHA-256+Salt加密 → '
         '调用UserDAO.updatePassword()更新数据库。')

add_para('')
add_para('UserServlet完整源代码：', bold=True)
add_code(user_servlet)

add_para('')
add_para('4.1.3 认证过滤器（AuthFilter.java）', bold=True)
add_para('AuthFilter通过@WebFilter("/*")注解拦截所有HTTP请求，实现统一的登录认证和权限校验。')
add_para('过滤逻辑分为三级：')
add_para('第一级——静态资源放行：路径以/css/、/js/、/images/、/uploads/开头的请求直接放行，不检查登录状态；')
add_para('第二级——公开路径放行：首页（/）、login.jsp、register.jsp、index.jsp、'
         '目的地列表、目的地详情、登录/注册处理、搜索接口等路径无需登录即可访问；')
add_para('第三级——受保护路径检查：其他所有路径需验证HttpSession中是否存在User对象 → '
         '若不存在则重定向到登录页 → 若路径包含/admin/还需额外检查user.isAdmin() → '
         '非管理员访问管理路径则重定向至首页。')

add_para('')
add_para('AuthFilter完整源代码：', bold=True)
add_code(auth_filter)

doc.add_page_break()

add_heading_styled('4.2 旅游目的地发布和编辑模块', level=2)
add_para('目的地管理（发布和编辑）是后台管理的核心功能，由DestinationServlet和DestinationService协同完成。')

add_para('')
add_para('4.2.1 目的地CRUD实现（DestinationServlet.java）', bold=True)
add_para('发布流程：管理员在destination_manage.jsp页面填写目的地信息表单'
         '（name/region/type/image/description/address/openTime/ticketPrice共8个字段） → '
         '表单通过POST提交到/admin/destination/*，action参数设为"create" → '
         'DestinationServlet.handleCreate()方法 → '
         '创建Destination对象并填充请求参数 → DestinationService.create()校验数据完整性 → '
         'DestinationDAO.create()生成INSERT SQL（使用PreparedStatement设置11个参数） → '
         '通过RETURN_GENERATED_KEYS获取自增主键ID → 重定向回管理页面，'
         '列表通过request.getSession().setAttribute()传递成功消息。')

add_para('编辑流程：管理员在管理页点击"编辑"按钮 → 页面展示预填充的目的地信息表单 → '
         '修改后提交POST（action="update"，附带隐藏id字段） → '
         'handleUpdate()先通过id查询现有记录（findById） → 更新字段值 → '
         'DestinationDAO.update()执行参数化UPDATE SQL，仅更新8个可变字段 → '
         '重定向回管理页面。')

add_para('删除流程：点击"删除"按钮 → JavaScript弹出confirm()确认对话框 → '
         '确认后POST提交（action="delete"，附带id） → '
         'handleDelete()调用DestinationDAO.delete(id) → 由于外键设置为ON DELETE CASCADE，'
         '该目的地关联的所有评论（comments）和投票（votes）记录自动级联删除。')

add_para('')
add_para('DestinationServlet完整源代码：', bold=True)
add_code(dest_servlet)

add_para('')
add_para('4.2.2 数据访问层查询实现（DestinationDAO.java核心方法）', bold=True)
add_para('DestinationDAO提供以下关键数据访问能力：')
add_para('（1）findById(int id)：通过子查询SELECT d.*, '
         '(SELECT COUNT(*) FROM votes WHERE destination_id = d.id) AS vote_count '
         'FROM destinations d WHERE d.id = ? 一次性获取目的地信息及其投票总数；')
add_para('（2）search(String keyword, String region, String type, String sortBy)：'
         '使用StringBuilder动态拼接WHERE条件，实现关键字模糊匹配（name LIKE ? OR description LIKE ?）、'
         '按region精确筛选、按type精确筛选，支持rating/popularity/newest三种排序方式；')
add_para('（3）updateRating(int destinationId)：执行UPDATE destinations d SET d.rating = '
         '(SELECT COALESCE(AVG(v.score), 0) FROM votes v WHERE v.destination_id = ?) '
         'WHERE d.id = ?，在每次评分操作后同步更新目的地综合评分；')
add_para('（4）incrementPopularity(int id)：执行UPDATE destinations SET popularity = popularity + 1 '
         'WHERE id = ?，用户每次访问详情页时自动增加目的地热度值。')

add_para('')
add_para('4.2.3 图片上传功能（ImageUploadServlet.java）', bold=True)
add_para('图片上传采用AJAX异步方式，避免整页刷新，其处理流程如下：')
add_para('（1）前端：管理员点击"Choose File"按钮选择本地图片 → JavaScript通过'
         'FileReader API读取文件进行本地预览显示 → 使用FormData + Fetch API将文件异步POST到'
         '/admin/upload-image → 收到服务器返回的JSON后自动将图片URL填入表单image输入框；')
add_para('（2）后端：ImageUploadServlet使用Apache Commons FileUpload库解析multipart/form-data请求 → '
         '校验文件扩展名（仅允许.jpg/.jpeg/.png/.gif/.webp/.bmp）→ '
         '校验文件大小（最大5MB = 5*1024*1024 bytes）→ '
         '通过UUID.randomUUID().toString()生成全局唯一文件名，保留原始扩展名 → '
         '将文件写入项目webapp/uploads/目录 → 返回JSON格式的成功响应（含生成的文件名和URL路径）。')

add_heading_styled('4.3 旅游浏览、评论与评分模块', level=2)
add_para('旅游目的地浏览、评论互动与星级评分是面向普通用户的核心前台功能，'
         '是用户体验的关键组成部分。')

add_para('')
add_para('4.3.1 列表浏览与搜索筛选', bold=True)
add_para('首页（index.jsp）：采用Bento Grid非对称12列网格布局展示最热门的6个目的地。'
         '页面顶部包含一个三合一的全局搜索栏（关键字输入框 + 地区下拉框 + 类型下拉框 + 搜索按钮），'
         '提交到/search接口。主体区域展示目的地卡片，每个卡片包含封面图片、类型标签（绿色tag）、'
         '地区信息、评分星级、投票数和票价信息。底部展示4列统计数据'
         '（目的地总数 / 地区数 / 类型数 / 总投票数），增强页面的信息密度和可信度。')

add_para('列表页（destination/list.jsp）：以表格视图展示全部目的地信息，包含缩略图、名称、'
         '地区/类型标签、评分、投票数、票价等列。顶部提供完整的筛选工具栏，'
         '支持地区下拉筛选、类型下拉筛选、关键字搜索和按评分/热度/最新排序。')

add_para('详情页（destination/detail.jsp）：左侧大面积展示目的地封面图（1200×600比例），'
         '下方依次为类型/地区标签、目的地名称（大标题）、描述文字。'
         '右侧以粘性定位（position: sticky; top: 80px）的信息卡片展示地址、开放时间、'
         '门票价格（¥0.00自动显示为"Free"）和访问热度。信息卡片下方为交互式评分区和评论区。')

add_para('')
add_para('4.3.2 交互式五星评分机制（核心互动模块）', bold=True)
add_para('评分模块采用纯JavaScript + AJAX实现，无需任何第三方库，工作流程如下：')

add_para('【前端实现（detail.jsp内嵌JavaScript）】')
add_para('（1）DOM结构：5个<button class="star-btn">按钮，通过data-star="1~5"属性标记星级；')
add_para('（2）鼠标悬停交互（mouseenter/mouseleave）：悬停时遍历所有star-btn，'
         '将≤当前星级的按钮添加.active类（颜色变为金黄色#e6a817），'
         '同步更新提示文字（很差/较差/一般/较好/非常好），移出后恢复当前实际评分状态；')
add_para('（3）点击提交（click → submitRating函数）：获取目标星级值 → '
         '构建URLSearchParams参数（destinationId + score + action=rate） → '
         '使用Fetch API发送POST请求到/vote → 解析返回的JSON → '
         '更新当前评分状态变量currentScore → 更新页面上的平均分数值、星级显示和投票人数 → '
         '显示Toast成功通知消息。同时实现了防重复提交（isSubmitting布尔锁）；')
add_para('（4）状态管理：页面加载时从服务端获取当前用户的评分状态（userScore和hasVoted），'
         '已评分的用户看到自己的星级高亮显示，提示文字显示"（点击可修改评分）"。')

add_para('')
add_para('【后端实现（VoteServlet + VoteService + VoteDAO）】')
add_para('VoteServlet设计为纯JSON API端点（不返回HTML页面），支持以下操作：')

add_para('• POST /vote?action=rate&destinationId=X&score=Y')
add_para('  处理逻辑：验证登录状态（未登录返回401） → 解析参数 → '
         'VoteService.rate(destId, userId, score)进行评分范围校验（clamp 1-5） → '
         'VoteDAO.rate()执行UPSERT逻辑（先findByUserAndDestination检查是否存在 → '
         '若存在则UPDATE已有记录的score，若不存在则INSERT新记录） → '
         '返回新平均分和投票总数 → 触发DestinationDAO.updateRating()同步更新destinations表rating字段。')

add_para('• POST /vote?action=unrate&destinationId=X')
add_para('  取消评分：VoteDAO.removeRating(destId, userId)执行DELETE → 重新计算平均分 → 返回结果。')

add_para('• GET /vote?destinationId=X')
add_para('  查询状态：返回该目的地的avgRating、voteCount、当前用户的hasVoted状态和userScore值，'
         '用于前端页面加载时的初始渲染。')

add_para('')
add_para('VoteServlet完整源代码（AJAX JSON接口）：', bold=True)
add_code(vote_servlet)

add_para('')
add_para('VoteDAO评分核心方法（UPSERT逻辑 + AVG计算）：', bold=True)
add_code(vote_dao)

add_para('')
add_para('VoteService业务规则校验（评分范围 + 评分同步）：', bold=True)
add_code(vote_service)

add_para('')
add_para('4.3.3 评论功能', bold=True)
add_para('评论功能通过CommentServlet + CommentService + CommentDAO三层协作实现：')
add_para('（1）发表评论：用户在详情页评论区输入文字内容，点击"提交评价"按钮 → '
         '表单POST到/comment，携带destinationId和content参数 → '
         'CommentServlet接收请求 → CommentService.addComment()校验内容非空 → '
         'CommentDAO.create()执行INSERT → 重定向回详情页，页面刷新后新评论显示在列表顶部；')
add_para('（2）评论列表展示：detail.jsp通过CommentService.getByDestination(destId)获取评论列表 → '
         'CommentDAO.findByDestinationId()使用JOIN查询'
         '（SELECT c.*, u.username FROM comments c JOIN users u ON c.user_id = u.id '
         'WHERE c.destination_id = ? ORDER BY c.created_at DESC）→ '
         '在JSP页面通过<% for (Comment c : comments) { %>循环渲染每条评论，'
         '显示评论者用户名首字母头像、用户名、评论时间和评论正文。')

add_para('')
add_para('SearchServlet搜索与筛选实现：', bold=True)
add_para('SearchServlet同时处理GET（显示搜索页）和POST（执行搜索）两种请求。'
         'POST处理时接收keyword/region/type/sortBy参数 → 调用DestinationDAO.search() → '
         '将结果列表存入request属性 → 转发到搜索结果JSP页面展示。'
         '支持无参数全量查询（即浏览全部目的地），同时支持任意组合条件筛选。')
add_code(search_servlet)

doc.add_page_break()

# ============================================================
# 5. 前台主要功能模块设计与实现
# ============================================================
add_heading_styled('5. 前台主要功能模块设计与实现', level=1)

add_para('前台页面基于Taste-Skill极简设计系统构建，所有页面共享一套CSS自定义属性（Custom Properties），'
         '通过设计令牌（Design Tokens）确保全局视觉风格的一致性。以下逐一说明各页面的设计与实现。')

add_para('')
add_para('5.1 全局导航栏（header.jsp）', bold=True)
add_para('导航栏采用Taste-Skill标志性的64px高度、粘性定位（position: sticky; top: 0）和1px底边框设计，'
         '是贯穿全站的可复用组件。左侧显示品牌Logo（菱形符号♦ + "TravelVote"文字），'
         '右侧动态显示导航链接，其内容根据用户登录状态和角色动态变化：')
add_para('• 未登录状态：显示"Destinations"浏览链接 + "Sign In"登录按钮；')
add_para('• 普通用户登录：显示"Destinations" + 用户名头像（圆形首字母Avatar） + "Log Out"退出按钮；')
add_para('• 管理员登录：额外显示"Admin"管理后台入口链接。')
add_para('导航栏下方为全局消息提示区域：'
         '从Session中读取error和success消息属性，用带颜色区分的alert组件展示（红色错误 / 绿色成功），'
         '展示后立即调用session.removeAttribute()清除，实现一次性消息提示。')
add_para('')
add_para('header.jsp完整代码：', bold=True)
add_code(header_jsp)

add_para('')
add_para('5.2 首页（index.jsp）', bold=True)
add_para('首页由四个功能区域组成：')
add_para('（1）Hero区域：大标题"Discover the world, one destination at a time." + '
         '副标题描述 + 两个CTA按钮（"Explore Now"主按钮跳转列表页，'
         '"Join Us"描边按钮跳转注册页），视觉上占据首屏主要面积；')
add_para('（2）全局搜索栏：三合一搜索组件（关键字输入框 + 地区下拉框 + 类型下拉框 + 搜索按钮），'
         '表单通过GET提交到/search接口，支持任意组合条件的即时搜索；')
add_para('（3）Bento Grid热门目的地：采用12列CSS Grid非对称布局，从Service层获取top 6目的地，'
         '每个卡片包含封面图（带onerror兜底降级）、类型标签、星级评分、名称、'
         '简介、地区/票价/投票数等元信息。卡片悬停时产生微妙的translateY(-2px)上浮效果和阴影增强；')
add_para('（4）统计数据区：4列统计卡片（Destinations总数 / Regions地区数 / Categories类型数 / '
         'Total Votes总投票数），通过IntersectionObserver驱动滚动渐入动画。')

add_para('')
add_para('5.3 详情页与评分组件（detail.jsp）', bold=True)
add_para('详情页是系统功能最密集的页面，集成了目的地信息展示、交互式五星评分和评论互动三大功能：')

add_para('评分区域设计：', bold=True)
add_para('• 评分概览区：大号数字显示当前平均评分（如"4.8"），配合5颗实心/空心星星直观展示评级，'
         '括号内显示评分人数；')
add_para('• 交互评分区（仅登录用户可见）：5颗可点击的星星按钮，支持hover预览效果'
         '（鼠标悬停到第3颗星时前3颗全部高亮变金黄#e6a817，并显示"一般"文字提示），'
         'click后触发AJAX调用submitRating()提交评分；')
add_para('• 登录提示区（未登录用户）：显示"登录后即可评分"链接，引导用户注册/登录。')

add_para('评论区域设计：', bold=True)
add_para('• 评论表单：包含一个textarea输入框和"提交评价"按钮，表单POST到/comment；')
add_para('• 评论列表：每条评论显示为独立的comment-item卡片，包含圆形首字母头像、用户名、'
         '评论时间（yyyy-MM-dd格式）和评论正文。空评论状态显示引导文案。')

add_para('')
add_para('详情页评分交互JavaScript代码（detail.jsp内嵌脚本，核心前端实现）：', bold=True)
add_code('''/* ===== Star Rating Interaction (detail.jsp) ===== */
(function() {
    var starsContainer = document.getElementById('stars-interactive');
    if (!starsContainer) return;
    var destId = starsContainer.dataset.destinationId;
    var currentScore = parseInt(starsContainer.dataset.userScore) || 0;
    var starBtns = starsContainer.querySelectorAll('.star-btn');
    var isSubmitting = false;

    // 高亮星星到第 N 颗
    function highlight(n) {
        starBtns.forEach(function(btn) {
            var star = parseInt(btn.dataset.star);
            btn.classList.toggle('active', star <= n);
        });
    }

    // 提交评分（AJAX Fetch POST）
    function submitRating(score) {
        if (isSubmitting) return;
        isSubmitting = true;
        fetch('<%=contextPath%>/vote', {
            method: 'POST',
            body: new URLSearchParams({
                destinationId: destId, score: score, action: 'rate'
            })
        }).then(function(res) { return res.json(); })
          .then(function(data) {
              isSubmitting = false;
              if (data.success) {
                  currentScore = score;
                  highlight(score);
                  // 实时更新平均分显示、星级、投票人数
                  showToast('评分成功！你给了 ' + score + ' 星');
              }
          }).catch(function(err) { isSubmitting = false; });
    }

    // 事件绑定：悬停预览 + 点击提交
    starBtns.forEach(function(btn) {
        btn.addEventListener('mouseenter', function() {
            highlight(parseInt(this.dataset.star));
        });
        btn.addEventListener('mouseleave', function() {
            highlight(currentScore);  // 恢复当前实际评分
        });
        btn.addEventListener('click', function() {
            submitRating(parseInt(this.dataset.star));
        });
    });
})();''')

add_para('')
add_para('CSS样式实现——星级评分组件样式（style.css）：', bold=True)
add_code('''/* ===== Star Rating Widget ===== */
.rating-section { padding: var(--space-5) 0; }
.rating-overview { display: flex; align-items: center; gap: var(--space-6); flex-wrap: wrap; }
.rating-number { font-size: 3rem; font-weight: 700; line-height: 1; }
.stars-display .star-display { font-size: 1.5rem; color: #DDD; }
.stars-display .star-display.filled { color: #e6a817; }
.star-btn {
    background: none; border: none; font-size: 2rem; cursor: pointer;
    color: #DDD; transition: color 0.15s, transform 0.15s; padding: 0 2px;
}
.star-btn:hover { transform: scale(1.15); }
.star-btn.active { color: #e6a817; }

/* ===== Toast Notification ===== */
.toast {
    position: fixed; bottom: 32px; left: 50%; transform: translateX(-50%);
    background: #111; color: #FFF; padding: 12px 24px; border-radius: 8px;
    font-size: 14px; z-index: 1000; opacity: 0; transition: opacity 0.3s;
}
.toast.show { opacity: 1; }''')

add_para('')
add_para('5.4 登录与注册页面', bold=True)
add_para('登录页（login.jsp）：中央居中卡片式布局，包含用户名/邮箱输入框（支持两种方式登录）、'
         '密码输入框和全宽"Sign In"按钮。卡片底部提供"Don\'t have an account? Sign up"跳转链接。'
         '表单POST到/user/login，由UserServlet.handleLogin()处理。')

add_para('登录页JSP代码：', bold=True)
add_code(login_jsp)

add_para('')
add_para('5.5 个人中心（user/profile.jsp）', bold=True)
add_para('个人中心提供用户资料编辑和密码修改两个功能区域。'
         '资料编辑表单预填充当前用户名和邮箱，修改后提交到/user/profile；'
         '密码修改表单需输入原密码、新密码和确认密码，前端校验两次输入一致性，'
         '后端验证原密码正确性后更新密码。')

add_para('')
add_para('5.6 Taste-Skill设计系统在前台的应用总结', bold=True)
add_para('Taste-Skill是一套基于CSS自定义属性的极简设计系统，'
         '本项目的所有前台页面均严格遵循其设计理念，具体应用包括：')

add_para('设计令牌（Design Tokens）：', bold=True)
add_code('''/* taste-tokens.css —— 全局CSS变量定义 */
:root {
    --color-canvas: #FBFBFA;              /* 页面背景 —— 暖白色 */
    --color-surface: #FFFFFF;              /* 卡片/容器背景 */
    --color-text: #111111;                 /* 主文字颜色 —— 近黑色 */
    --color-text-secondary: #555555;       /* 次要文字颜色 */
    --color-accent: #2D6A4F;              /* 强调色 —— 森林绿 */
    --color-border: #EAEAEA;              /* 边框 —— 1px浅灰 */
    --color-surface-hover: #F5F5F4;       /* 悬停态背景 */
    --radius-md: 8px;                     /* 统一圆角半径 */
    --shadow-sm: 0 1px 2px rgba(0,0,0,0.04); /* 卡片阴影 */
    --font-sans: 'Inter Tight', sans-serif;   /* 无衬线正文字体 */
    --font-serif: 'Newsreader', serif;         /* 衬线标题字体 */
    --space-4: 16px; --space-6: 24px; --space-8: 32px; /* 间距体系 */
}''')

add_para('')
add_para('系统支持的前端设计特性：')
add_bullet('滚动渐入动画：使用IntersectionObserver API监听元素进入视口，触发fade-up'
           '（opacity 0→1 + translateY 20px→0）过渡动画，营造优雅的加载体验；')
add_bullet('无障碍与偏好尊重：通过@media (prefers-reduced-motion: reduce)禁用动画，'
           '尊重用户的系统级减少动画偏好设置；')
add_bullet('暗色模式自适应：通过@media (prefers-color-scheme: dark)切换CSS变量值，'
           '自动适配系统暗色主题（背景变深色、文字变浅色、边框变暗）；')
add_bullet('响应式布局：768px断点以下将多列布局切换为单列，'
           'Bento Grid从12列降至1列，搜索栏从水平排列变为垂直堆叠，'
           '星级按钮尺寸缩小（2rem→1.5rem）。触控设备友好（按钮最小44px可点击区域）；')
add_bullet('微交互反馈：按钮点击时scale(0.98)下压效果（transition: transform 0.1s），'
           '卡片悬停时translateY(-2px)上浮效果和阴影增强，交互感生动自然。')

doc.add_page_break()

# ============================================================
# 6. 系统后台设计与实现
# ============================================================
add_heading_styled('6. 系统后台设计与实现', level=1)

add_heading_styled('6.1 数据库连接类设计与实现（DBUtil.java）', level=2)
add_para('DBUtil是整个系统的数据访问基础设施，封装了H2嵌入式数据库的连接管理、'
         '表结构自动创建和示例数据初始化三大职责。')

add_para('')
add_para('设计要点与关键技术决策：', bold=True)
add_para('（1）嵌入式数据库选型：选用H2 Database而非MySQL，原因有三：'
         '① 课程设计场景下零安装配置，评审老师无需搭建数据库环境；'
         '② 数据文件存储于项目./data/目录，可随项目源码一起分发；'
         '③ MODE=MySQL参数启用MySQL SQL语法兼容，未来迁移至MySQL仅需更换JDBC驱动和连接URL。')
add_para('（2）线程安全的懒初始化：使用volatile + synchronized双重检查锁定（Double-Checked Locking）模式，'
         '确保多线程环境下initDatabase()仅执行一次，且不会因锁竞争导致性能下降。')
add_para('（3）幂等初始化策略：initDatabase()方法开头执行DROP TABLE IF EXISTS（按外键依赖逆序：'
         'votes → comments → destinations → users），再执行CREATE TABLE，确保每次启动时的数据库状态一致。'
         '生产环境应改为CREATE TABLE IF NOT EXISTS。')
add_para('（4）资源安全管理：提供静态close(AutoCloseable... resources)工具方法，'
         '使用可变参数+try-catch静默关闭，简化DAO层代码中的资源释放逻辑。'
         'DAO层统一使用try-with-resources语法确保Connection/PreparedStatement/ResultSet自动关闭。')

add_para('')
add_para('DBUtil.java完整源代码（含12条示例数据种子）：', bold=True)
add_code(dbutil_code)

add_heading_styled('6.2 Servlet控制器类设计与实现', level=2)
add_para('系统共有6个Servlet类，统一采用Java Servlet 3.0+注解配置方式，避免在web.xml中'
         '编写冗长的<servlet>和<servlet-mapping>元素。')

add_para('')
add_para('Servlet路由映射总表：', bold=True)
add_code('''@WebServlet("/user/*")              → UserServlet
    POST /user/login              → 登录验证（用户名或邮箱 + 密码）
    POST /user/register           → 用户注册（用户名/邮箱/密码唯一性校验）
    POST /user/profile            → 个人资料编辑（用户名/邮箱更新）
    POST /user/change-password    → 密码修改（验证原密码 + 新密码加密）
    GET  /user/logout             → 安全退出（Session销毁）

@WebServlet("/admin/destination/*") → DestinationServlet
    POST action=create            → 创建新目的地
    POST action=update            → 编辑现有目的地
    POST action=delete            → 删除目的地（级联删除评论和投票）

@WebServlet("/admin/upload-image")  → ImageUploadServlet
    POST (multipart/form-data)    → AJAX图片上传（文件校验 + UUID命名 + 保存）

@WebServlet("/comment/*")           → CommentServlet
    POST                          → 发表评论（验证登录 + 校验内容非空）

@WebServlet("/vote")                → VoteServlet
    POST action=rate              → 提交/修改评分（AJAX JSON接口）
    POST action=unrate            → 取消评分（AJAX JSON接口）
    GET  ?destinationId=X         → 查询评分状态（AJAX JSON接口）

@WebServlet("/search")              → SearchServlet
    GET  (keyword/region/type/sortBy) → 搜索筛选目的地''')

add_para('')
add_para('每个Servlet遵循统一的设计模式：')
add_para('（1）doPost()/doGet()方法作为入口，首先验证用户登录状态和权限'
         '（管理员Servlet额外检查user.isAdmin()）；')
add_para('（2）从HttpServletRequest中提取action参数，通过if/else分支将请求分派到'
         '对应的私有处理方法（如handleLogin、handleCreate），保持方法职责单一；')
add_para('（3）调用Service层执行业务逻辑，Controller层不直接操作数据库；')
add_para('（4）根据请求类型决定响应方式：浏览器表单提交采用PRG模式（POST → Redirect → GET，'
         '防止表单重复提交），AJAX请求则设置Content-Type为application/json并返回JSON数据。')

add_heading_styled('6.3 登录验证处理页面', level=2)
add_para('登录页面（login.jsp）设计简洁高效，采用中央居中的卡片式布局。'
         '表单包含两个输入域（login支持用户名或邮箱、password密码）和一个全宽提交按钮。')

add_para('')
add_para('登录流程完整时序说明：')
add_para('（1）客户端：用户通过浏览器访问login.jsp → 浏览器渲染Taste-Skill风格登录卡片界面；')
add_para('（2）表单提交：用户填写登录凭证并点击"Sign In"按钮 → 浏览器以POST方式发送'
         'Form Data（login + password两个参数）到/user/login端点；')
add_para('（3）Filter链处理：AuthFilter检查路径/user/login是否在公开白名单中 → '
         '确认放行（登录处理路径无需预先登录）；')
add_para('（4）Servlet处理：UserServlet.doPost()接收请求 → 识别PathInfo为/login → '
         '调用handleLogin()方法；')
add_para('（5）Service层验证：UserService.login()方法中 → 调用UserDAO.findByUsernameOrEmail()'
         '按用户名或邮箱查询用户 → 若用户不存在返回null → 若用户存在，调用'
         'PasswordUtil.verify(inputPassword, storedPassword)验证密码哈希；')
add_para('（6）会话创建：验证通过 → 创建/获取HttpSession → '
         'session.setAttribute("user", user)将User对象存入会话 → '
         'session.setMaxInactiveInterval(30*60)设置30分钟超时；')
add_para('（7）路由跳转：根据user.getRole()判断角色 → 若role为admin → '
         '重定向到/admin/dashboard.jsp（管理员仪表盘）→ '
         '若role为user → 重定向到/（首页）；')
add_para('（8）失败处理：验证失败 → session.setAttribute("error", "Invalid username or password.") → '
         '重定向回login.jsp → header.jsp读取error属性并渲染红色错误提示。')

add_para('')
add_para('')
add_para('参考文献', bold=True)
add_para('[1] 李刚. 轻量级Java EE企业应用实战（第5版）[M]. 北京: 电子工业出版社, 2019.')
add_para('[2] 陈恒, 楼偶俊. Java Web开发实战（Spring+Spring MVC+MyBatis）[M]. 北京: 清华大学出版社, 2020.')
add_para('[3] 孙卫琴. Tomcat与Java Web开发技术详解（第3版）[M]. 北京: 电子工业出版社, 2019.')
add_para('[4] 王珊, 萨师煊. 数据库系统概论（第5版）[M]. 北京: 高等教育出版社, 2014.')
add_para('[5] H2 Database Engine Documentation [EB/OL]. https://www.h2database.com/html/main.html, 2024.')
add_para('[6] Apache Tomcat 9 Documentation [EB/OL]. https://tomcat.apache.org/tomcat-9.0-doc/, 2024.')
add_para('[7] Apache Commons FileUpload User Guide [EB/OL]. '
         'https://commons.apache.org/proper/commons-fileupload/using.html, 2024.')
add_para('[8] Taste-Skill Design System [EB/OL]. https://github.com/z75758/taste-skill, 2024.')
add_para('[9] Mozilla Developer Network. Using the Fetch API [EB/OL]. '
         'https://developer.mozilla.org/en-US/docs/Web/API/Fetch_API/Using_Fetch, 2024.')
add_para('[10] Java Servlet 3.0 Specification [S]. Oracle Corporation, 2009.')

# ── Save ──
output_path = r'F:\MyProject\TravelRating\report_section4_v2.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.0f} KB')
print('Done!')
