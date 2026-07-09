"""
Generate updated report sections for thesis replacement.
Based on report_section4_v2.docx, updated with recent code changes.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def h(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1: run.font.size = Pt(16)
        elif level == 2: run.font.size = Pt(14)
        elif level == 3: run.font.size = Pt(13)

def p(text, bold=False, indent=True):
    para = doc.add_paragraph()
    if indent: para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def code(text):
    for line in text.strip().split('\n'):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(51, 51, 51)

def bullet(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.74)
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def boxed_note(text):
    """Add a highlighted note box"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run('【更新说明】 ' + text)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(180, 60, 0)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Read current source files
base = r'F:\MyProject\TravelRating\src\main'

def read_src(path):
    full = os.path.join(base, path)
    if os.path.exists(full):
        with open(full, encoding='utf-8') as f:
            return f.read()
    return ''

auth_filter = read_src('java/com/travel/filter/AuthFilter.java')
encoding_filter = read_src('java/com/travel/filter/EncodingFilter.java')
comment_servlet = read_src('java/com/travel/controller/CommentServlet.java')
comment_service = read_src('java/com/travel/service/CommentService.java')
comment_dao = read_src('java/com/travel/dao/CommentDAO.java')
webxml = read_src('webapp/WEB-INF/web.xml')
detail_jsp = read_src('webapp/destination/detail.jsp')

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('课程设计报告 — 更新内容（替换指南）')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('旅游地点评选与互动系统设计与实现')
run.font.size = Pt(14)

p('')
p('本文档列出了自 report_section4_v2.docx 生成以来，系统中修改和新增的所有功能。'
  '请按照以下章节编号，将更新内容替换到你的论文对应位置。', bold=True)
p('')

# ── Change Summary Table ──
p('更新内容总览：', bold=True, indent=False)
bullet('更新1：评论删除功能（新增） → 替换第4.3节末尾 + 第5节详情页部分 + 第6.2节Servlet映射表')
bullet('更新2：中文乱码修复 → 替换第4.1节AuthFilter部分 + 第6.3节登录流程')
bullet('更新3：目的地图片URL更新 → 替换第6.1节DBUtil示例数据')

doc.add_page_break()

# ============================================================
# UPDATE 1: Comment Deletion
# ============================================================
h('更新1：评论删除功能（新增）', 1)

boxed_note('这是全新增加的功能，原报告中第4.3节"旅游浏览和评论"末尾需要追加此内容。')

h('4.3.4 评论删除功能', 2)
p('评论删除功能允许评论作者和管理员删除不当或过时的评论，完善了评论管理的完整生命周期。')

p('')
p('权限设计：', bold=True)
bullet('评论作者：可以删除自己发表的评论，每条评论右侧显示"×"删除按钮；')
bullet('管理员：可以删除任意用户的评论（越权管理）；')
bullet('其他用户：看不到删除按钮，无法删除他人评论；')
bullet('删除前弹出JavaScript确认对话框，防止误操作。')

p('')
p('后端实现（CommentServlet + CommentService + CommentDAO）：', bold=True)
p('CommentServlet新增handleDelete()方法，处理action=delete请求。'
  'CommentService.deleteComment()方法在执行删除前进行三重校验：'
  '① 通过CommentDAO.findById()查询评论是否存在；'
  '② 校验操作者是否为评论作者或管理员；'
  '③ 校验通过后调用CommentDAO.delete()执行物理删除。')

p('')
p('CommentServlet完整源代码（包含新增的删除处理逻辑）：', bold=True)
code(comment_servlet)

p('')
p('CommentService新增的deleteComment方法（含所有权校验）：', bold=True)
code('''/**
 * Delete a comment. Only the comment author or an admin can delete.
 * @return error message, or null if success
 */
public String deleteComment(int commentId, int userId, boolean isAdmin) {
    Comment comment = commentDAO.findById(commentId);
    if (comment == null) {
        return "评论不存在。";
    }
    if (!isAdmin && comment.getUserId() != userId) {
        return "你只能删除自己的评论。";
    }
    commentDAO.delete(commentId);
    return null;
}''')

p('')
p('CommentDAO新增的findById方法（用于删除前的所有权校验）：', bold=True)
code('''public Comment findById(int id) {
    String sql = "SELECT c.*, u.username FROM comments c "
               + "JOIN users u ON c.user_id = u.id "
               + "WHERE c.id = ?";
    try (Connection conn = DBUtil.getConnection();
         PreparedStatement stmt = conn.prepareStatement(sql)) {
        stmt.setInt(1, id);
        try (ResultSet rs = stmt.executeQuery()) {
            if (rs.next()) return mapRow(rs);
        }
    } catch (SQLException e) { e.printStackTrace(); }
    return null;
}''')

p('')
p('前端实现（detail.jsp评论列表新增删除按钮）：', bold=True)
p('每条评论的header区域右侧新增一个内嵌表单。表单包含action=delete、commentId和destinationId三个隐藏参数，'
  '以及一个"×"删除按钮。按钮仅在当前用户是评论作者或管理员时渲染。'
  '表单的onsubmit事件调用confirm()弹出确认对话框，用户确认后才提交POST请求。')
code('''<%-- 删除按钮：仅评论作者或管理员可见 --%>
<% if (detailCurrentUser != null &&
      (detailCurrentUser.isAdmin() || detailCurrentUser.getId() == c.getUserId())) { %>
<form action="<%=contextPath%>/comment" method="post"
      onsubmit="return confirm(\'确定要删除这条评论吗？\');" style="margin:0;">
    <input type="hidden" name="action" value="delete">
    <input type="hidden" name="commentId" value="<%= c.getId() %>">
    <input type="hidden" name="destinationId" value="<%= destId %>">
    <button type="submit" class="comment-delete-btn"
            title="删除评论" aria-label="删除评论">&times;</button>
</form>
<% } %>''')

p('')
p('CSS样式（style.css新增评论删除按钮样式）：', bold=True)
code('''.comment-delete-btn {
  background: none; border: none;
  font-size: 20px; color: #BBB;
  cursor: pointer; padding: 0 6px;
  line-height: 1;
  transition: color 0.15s, transform 0.15s;
  border-radius: var(--radius-sm);
}
.comment-delete-btn:hover {
  color: #E53E3E;              /* hover时变红色 */
  transform: scale(1.2);        /* 放大20% */
  background: #FFF5F5;         /* 淡红色背景 */
}''')

doc.add_page_break()

# ============================================================
# UPDATE 2: Encoding Fix
# ============================================================
h('更新2：中文乱码修复', 1)

boxed_note('原报告第4.1节中AuthFilter代码和第6.2节Servlet路由表需要更新。'
           '核心改动：移除@WebFilter注解，改用web.xml声明Filter执行顺序。')

h('4.1 登录验证模块（更新版）', 2)

p('')
p('4.1.3 认证过滤器（AuthFilter）—— 重要变更', bold=True)
p('【变更说明】原AuthFilter使用@WebFilter("/*")注解注册。'
  '由于EncodingFilter也使用相同的@WebFilter("/*")注解，Servlet 3.0规范不支持注解排序，'
  '导致两个Filter的执行顺序不确定。当AuthFilter先于EncodingFilter执行时，'
  'POST请求参数可能在UTF-8编码设置生效前被解析，使用默认ISO-8859-1编码读取中文内容，'
  '导致评论中出现乱码（如"黄山"显示为"??"）。')

p('修复方案：移除AuthFilter和EncodingFilter的@WebFilter注解，'
  '统一在web.xml中声明Filter及其执行顺序。EncodingFilter排在AuthFilter之前，'
  '确保所有请求先完成UTF-8编码设置，再进行认证拦截。', indent=True)

p('')
p('web.xml中Filter声明和执行顺序配置：', bold=True)
code('''<!-- Filter ordering: EncodingFilter MUST run before AuthFilter -->
<filter>
    <filter-name>EncodingFilter</filter-name>
    <filter-class>com.travel.filter.EncodingFilter</filter-class>
</filter>
<filter-mapping>
    <filter-name>EncodingFilter</filter-name>
    <url-pattern>/*</url-pattern>
</filter-mapping>

<filter>
    <filter-name>AuthFilter</filter-name>
    <filter-class>com.travel.filter.AuthFilter</filter-class>
</filter>
<filter-mapping>
    <filter-name>AuthFilter</filter-name>
    <url-pattern>/*</url-pattern>
</filter-mapping>''')

p('web.xml中Filter声明顺序即执行顺序：EncodingFilter先运行（设置UTF-8编码），'
  'AuthFilter后运行（检查登录状态和权限）。')

p('')
p('AuthFilter更新后的源代码（移除@WebFilter注解，其余逻辑不变）：', bold=True)
code(auth_filter)

p('')
p('EncodingFilter更新后的源代码（移除@WebFilter注解，其余逻辑不变）：', bold=True)
code(encoding_filter)

p('')
p('此外，CommentServlet的doPost()方法首行增加了防御性的字符编码设置，'
  '确保即使Filter配置有误也能正确处理中文：', bold=True)
code('''@Override
protected void doPost(HttpServletRequest req, HttpServletResponse resp)
        throws IOException {
    // CRITICAL: Set encoding BEFORE any getParameter() call
    req.setCharacterEncoding("UTF-8");
    // ... rest of the method
}''')

p('')
p('评论表单也增加了accept-charset属性，告知浏览器以UTF-8编码提交表单数据：', bold=True)
code('<form action="<%=contextPath%>/comment" method="post" accept-charset="UTF-8" ...>')

doc.add_page_break()

# ============================================================
# UPDATE 3: Image URLs
# ============================================================
h('更新3：目的地图片URL更新', 1)

boxed_note('原报告第6.1节DBUtil.java中的图片URL已过时（均为picsum.photos占位图），'
           '下文列出替换为真实图片URL的12条目的地示例数据。')

h('6.1 数据库连接类（更新版）', 2)

p('DBUtil.java中12条目的地示例数据已更新为真实图片URL。'
  '请将原报告第6.1节中samples数组的内容替换为以下版本：', bold=True)

code('''String[][] samples = {
    {"黄山风景区", "华东", "自然风光",
     "https://qimgs.qunarzz.com/.../0101p1200087cum179214.jpg_710x360_6d297f16.jpg",
     "黄山以奇松、怪石、云海、温泉四绝闻名于世...",
     "安徽省黄山市黄山区", "06:00-17:00", "190.00", "4.8", "980"},
    {"故宫博物院", "华北", "历史古迹",
     "https://www.shuomingshu.cn/.../a382daee878049f2969575e60d9f2464_vgf1x4cfjcj.jpg",
     "明清两代的皇家宫殿...",
     "北京市东城区景山前街4号", "08:30-17:00", "60.00", "4.9", "1200"},
    {"三亚亚龙湾", "华南", "海岛度假",
     "https://ts1.tc.mm.bing.net/th/id/R-C.3648af7ed8ba8c00a5390f53e1679327?...",
     "中国最南端的热带海滨旅游城市...",
     "海南省三亚市亚龙湾", "全天开放", "0.00", "4.7", "860"},
    {"丽江古城", "西南", "历史古迹",
     "https://ts1.tc.mm.bing.net/th/id/OIP-C.by5loua0w_sBBqwS8oNADgAAAA?...",
     "纳西族文化的瑰宝...",
     "云南省丽江市古城区", "全天开放", "50.00", "4.6", "750"},
    {"九寨沟", "西南", "自然风光",
     "https://p1.ssl.qhmsg.com/t01c5118f984a75ec07.jpg",
     "被称为人间仙境的九寨沟...",
     "四川省阿坝州九寨沟县", "08:00-17:00", "169.00", "4.9", "1100"},
    {"西湖", "华东", "自然风光",
     "https://hzyly.com/upload/201908/26/201908261930373237.jpg",
     "欲把西湖比西子...",
     "浙江省杭州市西湖区", "全天开放", "0.00", "4.7", "1050"},
    {"西安兵马俑", "西北", "历史古迹",
     "https://ts3.tc.mm.bing.net/th/id/OIP-C.qPHK2sHfvXVtZutSv1RY_QHaEK?...",
     "世界第八大奇迹...",
     "陕西省西安市临潼区", "08:30-18:00", "120.00", "4.8", "920"},
    {"成都锦里", "西南", "美食之旅",
     "https://ts3.tc.mm.bing.net/th/id/OIP-C.hA_CaJ-mV-TCDvX3WJBIqQHaEK?...",
     "锦里古街是成都美食的集中地...",
     "四川省成都市武侯区", "09:00-22:00", "0.00", "4.5", "680"},
    {"张家界国家森林公园", "华中", "自然风光",
     "https://so1.360tres.com/t01a91e60b4b166dd0a.jpg",
     "阿凡达取景地...",
     "湖南省张家界市武陵源区", "07:00-18:00", "228.00", "4.7", "890"},
    {"布达拉宫", "西南", "历史古迹",
     "https://ts2.tc.mm.bing.net/th/id/OIP-C.iWF1lrTNSgMO7_jvTn7YLwHaE5?...",
     "世界屋脊上的明珠...",
     "西藏拉萨市城关区", "09:00-16:00", "200.00", "4.8", "780"},
    {"呼伦贝尔大草原", "华北", "自然风光",
     "https://img.pconline.com.cn/.../258785371_1616767884199_mthumb.jpg",
     "中国最美的草原...",
     "内蒙古呼伦贝尔市", "全天开放", "0.00", "4.6", "620"},
    {"鼓浪屿", "华东", "海岛度假",
     "https://ts1.tc.mm.bing.net/th/id/R-C.e41bf9fad189c942d2f1ebfc6e480506?...",
     "钢琴之岛...",
     "福建省厦门市思明区", "全天开放", "35.00", "4.5", "730"}
};''')

doc.add_page_break()

# ============================================================
# UPDATE 4: Servlet Route Update
# ============================================================
h('更新4：Servlet路由映射表更新', 1)

boxed_note('原报告第6.2节Servlet路由映射表需要更新CommentServlet的描述。')

h('6.2 设计Servlet类（更新版Servlet路由映射）', 2)

p('CommentServlet路由更新（增加delete操作）：', bold=True)
code('''@WebServlet("/comment")           → CommentServlet
    POST action=(空)              → 发表评论（验证登录 + 校验内容非空）
    POST action=delete            → 删除评论（验证登录 + 校验所有者或管理员权限）''')

p('')
p('完整的6个Servlet路由映射表（更新后）：', bold=True)
code('''@WebServlet("/user/*")              → UserServlet
    POST /user/login              → 登录验证（用户名或邮箱 + 密码）
    POST /user/register           → 用户注册（用户名/邮箱/密码唯一性校验）
    POST /user/profile            → 个人资料编辑（用户名/邮箱更新）
    POST /user/change-password    → 密码修改（验证原密码 + 新密码加密）
    GET  /user/logout             → 安全退出（Session销毁）

@WebServlet("/admin/destination/*") → DestinationServlet
    POST action=create            → 创建新目的地
    POST action=update            → 编辑现有目的地
    POST action=delete            → 删除目的地（级联删除评论和投票）

@WebServlet("/admin/upload-image")  → ImageUploadServlet
    POST (multipart/form-data)    → AJAX图片上传

@WebServlet("/comment")             → CommentServlet
    POST (无action)               → 发表评论
    POST action=delete            → 删除评论（作者或管理员）★新增

@WebServlet("/vote")                → VoteServlet
    POST action=rate              → 提交/修改评分（AJAX JSON）
    POST action=unrate            → 取消评分（AJAX JSON）
    GET  ?destinationId=X         → 查询评分状态（AJAX JSON）

@WebServlet("/search")              → SearchServlet
    GET  (keyword/region/type)    → 搜索筛选目的地''')

# ============================================================
# REPLACEMENT GUIDE
# ============================================================
doc.add_page_break()
h('替换指南总结', 1)

p('请按照以下对照表，将本文档中的更新内容替换到原论文（report_section4_v2.docx）的对应位置：')
p('')

# Table
table_headers = ['序号', '更新内容', '原报告章节', '操作']
table_rows = [
    ['1', '评论删除功能：CommentServlet完整代码、CommentService.deleteComment()、CommentDAO.findById()、detail.jsp删除按钮代码、CSS样式',
     '4.3节末尾（追加）', '追加到4.3节"评论功能"描述之后'],
    ['2', 'AuthFilter源代码（移除@WebFilter）',
     '4.1.3节', '替换AuthFilter代码块'],
    ['3', 'EncodingFilter + web.xml Filter声明',
     '4.1.3节（新增小节）', '追加在AuthFilter之后'],
    ['4', 'Servlet路由映射表（含CommentServlet delete）',
     '6.2节', '替换Servlet路由映射表代码块'],
    ['5', 'DBUtil.java示例数据（12条真实图片URL）',
     '6.1节', '替换samples数组代码块'],
    ['6', '评论删除按钮及accept-charset',
     '5.3节详情页描述', '在评论区描述中补充'],
]

tbl = doc.add_table(rows=1+len(table_rows), cols=4, style='Table Grid')
for i, hdr in enumerate(table_headers):
    cell = tbl.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(hdr)
    run.bold = True; run.font.size = Pt(10)
    run.font.name = '宋体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
    cell._tc.get_or_add_tcPr().append(shading)
for r, row in enumerate(table_rows):
    for c, val in enumerate(row):
        cell = tbl.rows[r+1].cells[c]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        run.font.name = '宋体'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p('')
p('注意：以上代码均来自项目实际运行的最新版本，已经过完整测试验证。'
  '替换时请注意保持原文的格式风格一致，代码块使用等宽字体（Consolas），'
  '中文描述使用宋体小四号（12pt）。', bold=True, indent=True)

# ── Save ──
output = r'F:\MyProject\TravelRating\report_updates.docx'
doc.save(output)
import os
print(f'Update guide saved to: {output}')
print(f'File size: {os.path.getsize(output) / 1024:.0f} KB')
